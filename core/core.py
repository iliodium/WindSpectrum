import os
import random
import time
import glob
import pickle
from typing import Tuple, List
from multiprocessing import Process, Manager, managers
from concurrent.futures import ThreadPoolExecutor

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from scipy.signal import welch

from utils.utils import ks10, alpha_standards, wind_regions, speed_sp_b, interp_025_tpu, speed_sp_a, interp_016_tpu, \
    interp_016_real_tpu, interp_025_real_tpu
from utils.utils import interpolator as intp

import toml
import numpy as np
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scipy.interpolate import interp1d
import matplotlib.pyplot as plt

# Local imports
from plot.plot import Plot
from clipboard.clipboard import Clipboard
from utils.utils import (get_logger,
                         get_model_and_scale_factors,
                         rms,
                         rach,
                         obes_m,
                         obes_p,
                         converter_coordinates_to_real,
                         converter_coordinates,
                         generate_directory_for_report,
                         get_view_permutation_data,
                         get_sequence_permutation_data,
                         changer_sequence_coefficients,
                         get_base_angle,
                         to_dict,
                         to_multiprocessing_dict,
                         id_to_name,
                         get_model_and_scale_factors_interference,
                         )


class Core:
    config = toml.load('config.toml')

    _count_threads = config['core']['count_threads']  # количество запускаемых потоков при создании отчета

    def __init__(self, ex_clipboard=None):
        """Создание объекта для работы с буфером"""
        self.logger = get_logger('Core')
        self.logger.info('Создание ядра')
        self._manager = None
        self.clipboard_obj = Clipboard(ex_clipboard=ex_clipboard)
        self.logger.info('Ядро успешно создано')

    def save_clipboard(self):
        with open(f'{os.getcwd()}\\clipboard\\clipboard_binary', "wb") as clipboard_binary:
            dict_without_manager = to_dict(self.clipboard_obj.clipboard_dict)
            data_binary = pickle.dumps(dict_without_manager)
            clipboard_binary.write(data_binary)

    def get_plot_isofields(self,
                           db='isolated',
                           **kwargs):
        """Функция возвращает изополя"""
        if kwargs['pressure_plot_parameters']:
            type_plot = 'isofields_pressure'
        else:
            type_plot = 'isofields_coefficients'

        fig = self.clipboard_obj.get_isofields(db=db, type_plot=type_plot, **kwargs)

        return fig

    def get_plot_summary_spectres(self,
                                  db='interference',
                                  **kwargs):

        fig = self.clipboard_obj.get_summary_spectres(db=db, **kwargs)

        return fig

    def get_plot_summary_coefficients(self,
                                      db='isolated',
                                      **kwargs):

        fig = self.clipboard_obj.get_summary_coefficients(db, **kwargs)

        return fig

    def draw_welch_graphs(self, db, **kwargs):
        """Функция запускает отрисовку всех граффиков спектральной плотности мощности для выбранной модели.
        Отрисовка происходит в многопоточном режиме.
        """

        for mode in kwargs['mods']:
            args_welch_graphs = [(db, {'angle': str(angle), 'mode': mode}, kwargs) for angle in
                                 range(0, kwargs['angle_border'], 5)]
            for i in args_welch_graphs:
                self.welch_graphs_thread(i[0], **i[1], **i[2])

            with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                executor.map(lambda i: self.welch_graphs_thread(i[0], **i[1], **i[2]), args_welch_graphs)

    def welch_graphs_thread(self, db, **kwargs):
        """Функция для запуска потока генерации графиков спектральной плотности мощности и их сохранения."""
        mode = kwargs['mode']
        path_report = kwargs['path_report']
        model_scale = kwargs['model_scale']
        angle = int(kwargs['angle'])

        mode_fig = ' '.join(mode.lower().title().split('_'))
        path_folder = f'{path_report}\\Спектральная плотность мощности\\Логарифмическая шкала'
        if db == 'isolated':
            alpha = kwargs['alpha']
            file_name = f'Спектральная плотность мощности {model_scale} {alpha} {mode_fig} {angle:02}.png'

        elif db == 'interference':
            case = kwargs['case']
            file_name = f'Спектральная плотность мощности {model_scale} вариант {case} {mode_fig} {angle:02}.png'

        fig = self.get_plot_summary_spectres(db, **kwargs, scale='log', type_plot='summary_spectres')
        fig.savefig(f'{path_folder}\\{file_name}', bbox_inches='tight')
        plt.close(fig)

    def get_plot_pressure_tap_locations(self, db, **kwargs):
        return self.clipboard_obj.get_plot_pressure_tap_locations(db, **kwargs)

    def get_plot_model_3d(self, model_size):
        return self.clipboard_obj.get_plot_model_3d(model_size)

    def get_model_polar(self, model_size):
        return self.clipboard_obj.get_model_polar(model_size)

    def get_pseudocolor_coefficients(self, alpha, model_size, angle, mode):
        return self.clipboard_obj.get_pseudocolor_coefficients(alpha, model_size, angle, mode)

    def draw_summary_coefficients(self,
                                  db,
                                  **kwargs):

        """Функция запускает отрисовку всех графиков суммарных коэффициентов для выбранной модели.
        Отрисовка происходит в многопоточном режиме.
        """

        for mode in kwargs['mods']:
            args_summary_coefficients = [(db, {'angle': str(angle), 'mode': mode}, kwargs)
                                         for angle in range(0, kwargs['angle_border'], 5)]

            # for i in args_summary_coefficients:
            #     self.summary_coefficients_thread(i[0], **i[1], **i[2])

            with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                executor.map(lambda i: self.summary_coefficients_thread(i[0], **i[1], **i[2]),
                             args_summary_coefficients)

    def summary_coefficients_thread(self,
                                    db,
                                    **kwargs):
        """Функция для запуска потока генерации графиков суммарных аэродинамических коэффициентов модели
         и их сохранения."""
        angle = int(kwargs['angle'])
        mode = kwargs['mode']
        model_size = kwargs['model_size']
        path_report = kwargs['path_report']

        mode_fig = ' '.join(mode.lower().title().split('_'))
        if db == 'isolated':
            alpha = kwargs['alpha']
            file_name = f'Суммарные аэродинамические коэффициенты {mode_fig} ' \
                        f'{" ".join(model_size)} {alpha} {angle:02}.png'
        elif db == 'interference':
            case = kwargs['case']
            file_name = f'Суммарные аэродинамические коэффициенты {mode_fig} ' \
                        f'{" ".join(model_size)} вариант {case} {angle:02}.png'
        path_sum = f'{path_report}\\Суммарные аэродинамические коэффициенты\\Декартовая система координат\\' \
                   f'{angle:02}'
        if not os.path.isdir(path_sum):
            os.mkdir(path_sum)

        fig = self.get_plot_summary_coefficients(db, **kwargs)

        fig.savefig(f'{path_sum}\\{file_name}', bbox_inches='tight')
        plt.close(fig)

    def isofields_coefficients_thread(self,
                                      db,
                                      **kwargs):
        """Функция для запуска потока генерации изополей и их сохранения."""
        model_size = kwargs['model_size']
        angle = kwargs['angle']
        mode = kwargs['mode']
        path_report = kwargs['path_report']
        if db == 'isolated':
            alpha = kwargs['alpha']
            file_name = f'Изополя {" ".join(model_size)} {alpha} {int(angle):02} {mode}.png'

        elif db == 'interference':
            case = kwargs['case']
            file_name = f'Изополя {" ".join(model_size)} вариант {case} {int(angle):02} {mode}.png'

        fig = self.get_plot_isofields(db, pressure_plot_parameters=None, **kwargs)

        fig.savefig(
            f'{path_report}\\Изополя ветровых нагрузок и воздействий\\Коэффициенты\\{mode}\\{file_name}',
            bbox_inches='tight')

        plt.close(fig)

    def draw_isofields_coefficients(self,
                                    db,
                                    **kwargs, ):
        """Функция запускает отрисовку всех изополей выбранной модели.
        Отрисовка происходит в многопоточном режиме.
        """
        mods = kwargs['mods']
        angle_border = kwargs['angle_border']

        # if db == 'isolated':
        #     model_size = kwargs['model_size']
        #     alpha = kwargs['alpha']

        # self.logger.info(f'Отрисовка изополей размеры = {" ".join(model_size)} альфа = {alpha}')

        for mode in mods:
            args = [(db, {'angle': angle, 'mode': mode}, kwargs) for angle in range(0, angle_border, 5)]

            # for i in args:
            #     self.isofields_coefficients_thread(i[0], **i[1], **i[2])

            with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                executor.map(lambda i: self.isofields_coefficients_thread(i[0], **i[1], **i[2]), args)

            # self.logger.info(f'Отрисовка изополей размеры = {" ".join(model_size)} альфа = {alpha} завершена')
        # elif db == 'interference':
        #     for mode in mods:
        #         args = [(alpha, model_size, str(angle), mode, path_report) for angle in
        #                 range(0, angle_border, 5)]
        #
        #         with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
        #             executor.map(lambda i: self.isofields_coefficients_thread(*i), args)

    def isofields_pressure_thread(self,
                                  db,
                                  **kwargs):
        model_size = kwargs['model_size']
        angle = kwargs['angle']
        mode = kwargs['mode']
        path_report = kwargs['path_report']

        if db == 'isolated':
            alpha = kwargs['alpha']
            file_name = f'Изополя давления {" ".join(model_size)} {alpha} {int(angle):02} {mode}.png'

        elif db == 'interference':
            case = kwargs['case']
            file_name = f'Изополя давления {" ".join(model_size)} вариант {case} {int(angle):02} {mode}.png'

        fig = self.get_plot_isofields(db, **kwargs)

        fig.savefig(
            f'{path_report}\\Изополя ветровых нагрузок и воздействий\\Давление\\{mode}\\{file_name}',
            bbox_inches='tight')
        plt.close(fig)

    def draw_isofields_pressure(self, db, **kwargs):
        mods = kwargs['mods']
        angle_border = kwargs['angle_border']

        for mode in mods:
            args = [(db, {'angle': angle, 'mode': mode}, kwargs) for angle in range(0, angle_border, 5)]

            with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                executor.map(lambda i: self.isofields_pressure_thread(i[0], **i[1], **i[2]), args)

    def draw_pseudocolor_coefficients(self, alpha, model_size, angle_border, path_report, mods):
        for mode in mods:
            args = [(alpha, model_size, str(angle), mode, path_report) for angle in
                    range(0, angle_border, 5)]
            # for i in args:
            #     self.pseudocolor_coefficients_thread(*i)
            with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                executor.map(lambda i: self.pseudocolor_coefficients_thread(*i), args)

    def pseudocolor_coefficients_thread(self, alpha, model_size, angle, mode, path_report):
        file_name = f'Мозаика коэффициентов {" ".join(model_size)} {alpha} {int(angle):02} {mode}.png'

        fig = self.get_pseudocolor_coefficients(alpha, model_size, angle, mode)

        fig.savefig(
            f'{path_report}\\Мозаика коэффициентов\\{mode}\\{file_name}', bbox_inches='tight')
        plt.close(fig)

    def draw_summary_coefficients_polar(self,
                                        db,
                                        **kwargs):
        """Функция запускает отрисовку всех графиков суммарных коэффициентов в полярной системе для выбранной модели.
        Отрисовка происходит в многопоточном режиме.
        """

        args = [(db, {'mode': mode}, kwargs) for mode in kwargs["mods"]]

        with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
            polar_plots = list(
                executor.map(lambda i: self.get_plot_summary_coefficients_polar(i[0], **i[1], **i[2]), args))
        path_folder = f'{kwargs["path_report"]}\\Суммарные аэродинамические коэффициенты\\Полярная система координат'

        for mode, fig in zip(kwargs["mods"], polar_plots):
            file_name = f'Суммарные аэродинамические коэффициенты Cx Cy CMz {" ".join(kwargs["model_size"])} {id_to_name[mode]}.png'
            fig.savefig(f'{path_folder}\\{file_name}', bbox_inches='tight')
            plt.close(fig)

    def get_plot_summary_coefficients_polar(self,
                                            db='isolated',
                                            **kwargs):
        """Функция возвращает графики суммарных коэффициентов в полярной системе координат.
        Если графики суммарных коэффициентов в полярной системе координат отсутствуют в буфере, запускается отрисовка.
        """
        if db == 'isolated':
            alpha = kwargs['alpha']
            mode = kwargs['mode']
            model_size = kwargs['model_size']
            model_scale = kwargs['model_scale']
            angle_border = kwargs['angle_border']

            self.logger.info(f'Запрос суммарных коэффициентов в '
                             f'полярной системе координат размеры = {" ".join(list(model_size))} '
                             f'альфа = {alpha} режим = {mode.ljust(4, " ")} из буфера')

            mods = {
                'mean': np.mean,
                'rms': rms,
                'std': np.std,
                'max': np.max,
                'min': np.min,
                'Расчетное': rach,
                'rach': rach,
                'obesP': obes_p,
                'obesM': obes_m,
                'Обеспеченность +': obes_m,
                'Обеспеченность -': obes_p
            }

            id_fig = f'summary_coefficients_Cx_Cy_CMz_polar_{"_".join(model_size)}_{mode}'

            if not self.clipboard_obj.clipboard_dict['isolated'][alpha][model_scale]['const_stuff'].get(id_fig):
                self.clipboard_obj.get_coordinates(db='isolated', alpha=alpha, model_scale=model_scale)

                args_cmz = [('isolated', {'alpha': alpha, 'model_scale': model_scale, 'angle': str(angle)}) for angle in
                            range(0, angle_border, 5)]
                with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                    list_cmz = list(executor.map(lambda i: self.clipboard_obj.get_cmz(i[0], **i[1]), args_cmz))

                args_cx_cy = [('isolated', {'alpha': alpha, 'model_scale': model_scale, 'angle': str(angle)}) for angle
                              in range(0, angle_border, 5)]
                with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                    cx_cy = np.array(
                        list(executor.map(lambda i: self.clipboard_obj.get_cx_cy(i[0], **i[1]), args_cx_cy)))

                list_cx = cx_cy[:, 0]
                list_cy = cx_cy[:, 1]

                list_norm_cmz = list(map(mods[mode], list_cmz))
                list_norm_cx = list(map(mods[mode], list_cx))
                list_norm_cy = list(map(mods[mode], list_cy))

                x_scale, y_scale = Plot.scaling_data(list_norm_cx, list_norm_cy, angle_border=angle_border)
                cmz_scale = Plot.scaling_data(list_norm_cmz, angle_border=angle_border)

                data = {
                    'Cx': x_scale,
                    'Cy': y_scale,
                    'CMz': cmz_scale,
                }

                fig = Plot.polar_plot(db, data=data, title=mode, model_size=model_size, alpha=alpha)
                self.clipboard_obj.clipboard_dict['isolated'][alpha][model_scale]['const_stuff'][id_fig] = fig

            fig = self.clipboard_obj.clipboard_dict['isolated'][alpha][model_scale]['const_stuff'].get(id_fig)

            self.logger.info(f'Запрос суммарных коэффициентов в полярной системе координат размеры = '
                             f'{" ".join(list(model_size))} альфа = {alpha} режим = {mode.ljust(4, " ")} из буфера успешно выполнен')
        elif db == 'interference':
            case = kwargs['case']
            mode = kwargs['mode']
            model_size = kwargs['model_size']
            model_scale = kwargs['model_scale']

            mods = {
                'mean': np.mean,
                'rms': rms,
                'std': np.std,
                'max': np.max,
                'min': np.min,
                'Расчетное': rach,
                'rach': rach,
                'obesP': obes_p,
                'obesM': obes_m,
                'Обеспеченность +': obes_m,
                'Обеспеченность -': obes_p
            }

            id_fig = f'summary_coefficients_Cx_Cy_CMz_polar_{"_".join(model_size)}_{mode}_{db}_{case}'

            if not self.clipboard_obj.clipboard_dict['interference'][model_scale][case]['const_stuff'].get(id_fig):
                self.clipboard_obj.get_coordinates(db='interference', case=case, model_scale=model_scale)
                args_cmz = [('interference', {'case': case, 'model_scale': model_scale, 'angle': angle}) for angle
                            in
                            range(0, 360, 5)]
                with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                    list_cmz = list(executor.map(lambda i: self.clipboard_obj.get_cmz(i[0], **i[1]), args_cmz))

                args_cx_cy = [('interference', {'case': case, 'model_scale': model_scale, 'angle': angle}) for
                              angle in
                              range(0, 360, 5)]
                with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                    cx_cy = np.array(
                        list(executor.map(lambda i: self.clipboard_obj.get_cx_cy(i[0], **i[1]), args_cx_cy)))

                list_cx = cx_cy[:, 0]
                list_cy = cx_cy[:, 1]

                list_norm_cmz = list(map(mods[mode], list_cmz))
                list_norm_cx = list(map(mods[mode], list_cx))
                list_norm_cy = list(map(mods[mode], list_cy))

                list_norm_cmz.append(list_norm_cmz[0])
                list_norm_cx.append(list_norm_cx[0])
                list_norm_cy.append(list_norm_cy[0])

                data = {
                    'Cx': list_norm_cx,
                    'Cy': list_norm_cy,
                    'CMz': list_norm_cmz,
                }

                fig = Plot.polar_plot(db, data=data, title=mode, model_size=model_size, case=case)
                self.clipboard_obj.clipboard_dict['interference'][model_scale][case]['const_stuff'][id_fig] = fig

            fig = self.clipboard_obj.clipboard_dict['interference'][model_scale][case]['const_stuff'].get(id_fig)

        return fig

    def get_envelopes(self,
                      db='isolated',
                      **kwargs
                      ):
        """Функция возвращает графики огибающих.
        Если огибающие отсутствуют в буфере, запускается отрисовка.
        """
        mods = ['mean', 'rms', 'std', 'max', 'min']

        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale = kwargs['model_scale']
            angle = kwargs['angle']
            # mods = kwargs['mods']
            self.logger.info(f'Запрос огибающих параметры = {" ".join(list(model_scale))} '
                             f'альфа = {alpha} угол = {angle.rjust(2, "0")} из буфера')

            if not self.clipboard_obj.clipboard_dict['isolated'][alpha][model_scale][angle].get('envelopes'):
                self.logger.info(f'Отрисовка огибающих {" ".join(list(model_scale))} {alpha} {int(angle):02}')

                pressure_coefficients = self.clipboard_obj.get_pressure_coefficients('isolated', alpha=alpha,
                                                                                     model_name=model_scale,
                                                                                     angle=angle)
                figs = Plot.envelopes(pressure_coefficients, alpha, model_scale, angle, mods)

                for ind, val in enumerate(figs):
                    self.clipboard_obj.clipboard_dict['isolated'][alpha][model_scale][angle][
                        f'envelopes_{ind}_{"_".join(mods)}'] = val
                self.clipboard_obj.clipboard_dict['isolated'][alpha][model_scale][angle]['envelopes'] = True
                self.logger.info(f'Отрисовка огибающих {" ".join(list(model_scale))} {alpha} {int(angle):02} завершена')

            figs = []
            i = 0
            while True:
                if not self.clipboard_obj.clipboard_dict['isolated'][alpha][model_scale][angle].get(
                        f'envelopes_{i}_{"_".join(mods)}'):
                    break
                else:
                    figs.append(
                        self.clipboard_obj.clipboard_dict['isolated'][alpha][model_scale][angle][
                            f'envelopes_{i}_{"_".join(mods)}'])
                    i += 1

            self.logger.info(f'Запрос огибающих параметры = {" ".join(list(model_scale))} '
                             f'альфа = {alpha} угол = {angle.rjust(2, "0")} из буфера успешно выполнен')

        elif db == 'interference':
            case = kwargs['case']
            model_scale = kwargs['model_scale']
            angle = int(kwargs['angle'])
            if not self.clipboard_obj.clipboard_dict['interference'][model_scale][case][angle].get('envelopes'):
                pressure_coefficients = self.clipboard_obj.get_pressure_coefficients('interference', case=case,
                                                                                     model_name=model_scale,
                                                                                     angle=angle)

                figs = Plot.envelopes(pressure_coefficients, case, model_scale, angle, mods)

                for ind, val in enumerate(figs):
                    self.clipboard_obj.clipboard_dict['interference'][model_scale][case][angle][
                        f'envelopes_{ind}_{"_".join(mods)}'] = val
                self.clipboard_obj.clipboard_dict['interference'][model_scale][case][angle]['envelopes'] = True

            figs = []
            i = 0
            while True:
                if not self.clipboard_obj.clipboard_dict['interference'][model_scale][case][angle].get(
                        f'envelopes_{i}_{"_".join(mods)}'):
                    break
                else:
                    figs.append(
                        self.clipboard_obj.clipboard_dict['interference'][model_scale][case][angle][
                            f'envelopes_{i}_{"_".join(mods)}'])
                    i += 1

        return figs

    def draw_envelopes(self,
                       db, **kwargs):
        """Функция запускает отрисовку всех огибающих выбранной модели.
        Отрисовка происходит в многопоточном режиме.
        """
        args = [(db, {'angle': str(angle)}, kwargs) for angle in range(0, kwargs['angle_border'], 5)]

        # for i in args:
        #     self.envelopes_thread(i[0], **i[1], **i[2])

        # Для мощных пк, если есть 10ГБ оперативной памяти лишней
        with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
            executor.map(lambda i: self.envelopes_thread(i[0], **i[1], **i[2]), args)

    def envelopes_thread(self,
                         db,
                         **kwargs):
        """Функция для запуска потока генерации графиков огибающих модели и их сохранения."""
        figs = self.get_envelopes(db=db, **kwargs)
        path_report = kwargs['path_report']
        model_scale = kwargs['model_scale']
        angle = kwargs['angle']
        mods = kwargs['mods']

        if db == 'isolated':
            alpha = kwargs['alpha']
            path_envelopes = f'{path_report}\\Огибающие\\{model_scale} {alpha} {angle}'

        elif db == 'interference':
            case = kwargs['case']
            path_envelopes = f'{path_report}\\Огибающие\\{model_scale} {case} {angle}'

        if not os.path.isdir(path_envelopes):
            os.mkdir(path_envelopes)

        for i in range(len(figs)):
            if db == 'isolated':
                dbb = alpha
            elif db == 'interference':
                dbb = case

            file_name = f'Огибающие {model_scale} {dbb} {int(angle):02} ' \
                        f'{i * 100} - {(i + 1) * 100} {" ".join(mods)}.png'
            figs[i].savefig(f'{path_envelopes}\\{file_name}', bbox_inches='tight')
            plt.close(figs[i])

    @staticmethod
    def check_alive_proc(proc, button, spinner):
        while proc.is_alive():
            time.sleep(1)

        spinner.active = False
        button.disabled = False

    def wrapper_for_clipboard(self):
        if not self._manager:
            self._manager = Manager()

        self.clipboard_obj.clipboard_dict = to_multiprocessing_dict(self.clipboard_obj.clipboard_dict, self._manager)

    def draw_plot_model_3d(self, model_size, path_report):
        fig = self.get_plot_model_3d(model_size)
        fig.savefig(f'{path_report}\\Модель\\Модель трехмерная.png', bbox_inches='tight')
        plt.close(fig)

    def draw_model_polar(self, model_size, path_report):
        fig = self.get_model_polar(model_size)
        fig.savefig(f'{path_report}\\Модель\\Модель в полярной системе.png', bbox_inches='tight')
        plt.close(fig)

    def draw_plot_pressure_tap_locations(self, db, **kwargs):
        path_report = kwargs['path_report']
        alpha = kwargs['alpha']
        model_size = kwargs['model_size']
        fig = self.get_plot_pressure_tap_locations(db, alpha=alpha, model_size=model_size)

        fig.savefig(f'{path_report}\\Модель\\Развертка модели.png', bbox_inches='tight')
        plt.close(fig)

    def get_summary_coefficients_statistics(self, db, **kwargs):
        mods = kwargs['mods']

        accuracy_values = 2
        statistics = []

        data = dict()

        cmz = self.clipboard_obj.get_cmz(db, **kwargs)
        data['cmz'] = cmz

        cx, cy = self.clipboard_obj.get_cx_cy(db, **kwargs)
        data['cx'] = cx
        data['cy'] = cy

        functions = {
            'max': lambda d: np.max(d, axis=0).round(accuracy_values),
            'mean': lambda d: np.mean(d, axis=0).round(accuracy_values),
            'min': lambda d: np.min(d, axis=0).round(accuracy_values),
            'std': lambda d: np.std(d, axis=0).round(accuracy_values),
            'rms': lambda d: rms(d).round(accuracy_values),
            'rach': lambda d: rach(d).round(accuracy_values),
            'obesP': lambda d: obes_p(d).round(accuracy_values),
            'obesM': lambda d: obes_m(d).round(accuracy_values)
        }

        for name in data.keys():
            statistics_name = []
            statistics_name.append(name.capitalize())

            for mode in mods:
                temp_val = functions[mode](data[name])
                statistics_name.append(int(temp_val) if temp_val.is_integer() else f'{temp_val:.2f}')

            statistics.append(statistics_name)

        return statistics

    def get_sensor_statistics(self, db, **kwargs):
        model_size = kwargs['model_size']
        mods = kwargs['mods']

        accuracy_values = 2

        breadth_real, depth_real, height_real = float(model_size[0]), float(model_size[1]), float(model_size[2])
        x, z = self.clipboard_obj.get_coordinates(db, **kwargs)
        sensors_on_model = len(x)

        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale = kwargs['model_scale']
            face_number = self.clipboard_obj.get_face_number(alpha, model_scale)
            kwargs['angle'] = str(kwargs['angle'])

        elif db == 'interference':
            kwargs['angle'] = int(kwargs['angle'])
            face_number = []
            for _ in range(sensors_on_model // 28):
                face_number.extend([1 for _ in range(7)])
                face_number.extend([2 for _ in range(7)])
                face_number.extend([3 for _ in range(7)])
                face_number.extend([4 for _ in range(7)])

        kwargs['model_name'] = kwargs['model_scale']

        x_new, y_new = converter_coordinates(x, breadth_real, depth_real, face_number, sensors_on_model)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients(db, **kwargs)

        pressure_coefficients_t = pressure_coefficients.T

        functions = {'x': lambda: x_new,
                     'y': lambda: y_new,
                     'z': lambda: np.array(z).round(1),
                     'max': lambda: np.max(pressure_coefficients, axis=0).round(accuracy_values),
                     'mean': lambda: np.mean(pressure_coefficients, axis=0).round(accuracy_values),
                     'min': lambda: np.min(pressure_coefficients, axis=0).round(accuracy_values),
                     'std': lambda: np.std(pressure_coefficients, axis=0).round(accuracy_values),
                     'rms': lambda: [rms(i).round(accuracy_values) for i in pressure_coefficients_t],
                     'rach': lambda: [rach(i).round(accuracy_values) for i in pressure_coefficients_t],
                     'obesP': lambda: [obes_p(i).round(accuracy_values) for i in pressure_coefficients_t],
                     'obesM': lambda: [obes_m(i).round(accuracy_values) for i in pressure_coefficients_t]}

        statistics_of_angle = [[i for i in range(1, sensors_on_model + 1)]]

        for k in mods:
            statistics_of_angle.append(functions[k]())

        return np.array(statistics_of_angle).T

    def preparation_for_report(self,
                               pressure_plot_parameters,
                               content,
                               button,
                               spinner,
                               **kwargs):
        # if isinstance(self.clipboard_obj.clipboard_dict, dict):
        #     self.wrapper_for_clipboard()

        args = (self.clipboard_obj.clipboard_dict, pressure_plot_parameters, content)

        Core.start_report(*args, **kwargs)

        # proc = Process(target=Core.start_report, args=args, kwargs=kwargs)
        # proc.start()

        # executor = ThreadPoolExecutor(max_workers=1)
        # executor.map(lambda i: Core.check_alive_proc(*i), ((proc, button, spinner),))

    @staticmethod
    def start_report(clipboard_dict, pressure_plot_parameters, content, **kwargs):
        new_core = Core(clipboard_dict)
        new_core.report(pressure_plot_parameters, content, **kwargs)

    def report(self, pressure_plot_parameters, content, **kwargs):
        t1 = time.time()
        db = kwargs['db']
        del kwargs['db']
        model_size = kwargs['model_size']
        breadth, depth, height = model_size

        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale, scale_factors = get_model_and_scale_factors(*model_size, alpha)

            name_report = f'Отчет ширина {breadth} глубина {depth} высота {height} альфа {alpha}'
            if model_scale[0] == model_scale[1]:
                angle_border = 50
            else:
                angle_border = 95

        elif db == 'interference':
            case = kwargs['case']
            model_scale, scale_factors = get_model_and_scale_factors_interference(*model_size)

            name_report = f'Отчет ширина {breadth} глубина {depth} высота {height} вариант {case}'
            angle_border = 360

        mode_names = {
            'mean': 'средних',
            'std': 'стандартных отклонений',
            'max': 'максимальных',
            'min': 'минимальных',
        }

        # angle_border = 10

        current_path = os.getcwd()

        path_report = f'{current_path}\\Отчеты\\{name_report}'

        kwargs['model_scale'] = model_scale
        kwargs['angle_border'] = angle_border
        kwargs['path_report'] = path_report

        generate_directory_for_report(current_path, name_report)

        # Только эти графики используют координаты
        if any((content['isofieldsCoefficients'][0],
                # content['pressureTapLocations'][0] if db == 'isolated' else None,
                content['pseudocolorCoefficients'][0],
                content['isofieldsPressure'][0],
                content['summaryCoefficients'][0],
                content['statisticsSensors'][0],
                # content['statisticsSummaryCoefficients'][0],
                content['summarySpectres'][0])):
            self.clipboard_obj.get_coordinates(db=db, **kwargs)
            args = [(db, {'angle': angle}, kwargs) for angle in range(0, angle_border, 5)]

            with ThreadPoolExecutor(max_workers=Core._count_threads) as executor:
                executor.map(lambda i: self.clipboard_obj.get_pressure_coefficients(i[0], **i[1], **i[2]), args)

        # Отрисовка графиков
        self.logger.info(f'Отрисовка графиков')
        # isofieldsPressure
        if content['isofieldsPressure'][0]:
            mods_isofieldsPressure = [k for k in content['isofieldsPressure'][1].keys() if
                                      content['isofieldsPressure'][1][k]]
            self.draw_isofields_pressure(db, **kwargs, mods=mods_isofieldsPressure,
                                         pressure_plot_parameters=pressure_plot_parameters)
            self.logger.info(f'isofieldsPressure')

        # isofieldsCoefficients
        if content['isofieldsCoefficients'][0]:
            mods_isofieldsCoefficients = [k for k in content['isofieldsCoefficients'][1].keys() if
                                          content['isofieldsCoefficients'][1][k]]
            self.draw_isofields_coefficients(db, **kwargs, mods=mods_isofieldsCoefficients)
            self.logger.info(f'isofieldsCoefficients')

        # pseudocolorCoefficients
        if content['pseudocolorCoefficients'][0] and db == 'isolated':
            mods_pseudocolorCoefficients = [k for k in content['pseudocolorCoefficients'][1].keys() if
                                            content['pseudocolorCoefficients'][1][k]]
            self.draw_pseudocolor_coefficients(alpha=kwargs['alpha'], model_size=kwargs['model_size'],
                                               angle_border=kwargs['angle_border'], path_report=kwargs['path_report'],
                                               mods=mods_pseudocolorCoefficients)
            self.logger.info(f'pseudocolorCoefficients')
        # envelopes
        if content['envelopes'][0]:
            mods_envelopes = [k for k in content['envelopes'][1].keys() if content['envelopes'][1][k]]
            self.draw_envelopes(db, **kwargs, mods=mods_envelopes)
            self.logger.info(f'envelopes')

        # polarSummaryCoefficients
        if content['polarSummaryCoefficients'][0]:
            mods_polarSummaryCoefficients = [k for k in content['polarSummaryCoefficients'][1].keys() if
                                             content['polarSummaryCoefficients'][1][k]]
            self.draw_summary_coefficients_polar(db, **kwargs, mods=mods_polarSummaryCoefficients)
            self.logger.info(f'polarSummaryCoefficients')

        # summaryCoefficients
        if content['summaryCoefficients'][0]:
            mods_summaryCoefficients = [k for k in content['summaryCoefficients'][1].keys() if
                                        content['summaryCoefficients'][1][k]]
            self.draw_summary_coefficients(db, **kwargs, mods=mods_summaryCoefficients)
            self.logger.info(f'summaryCoefficients')

        # summarySpectres
        if content['summarySpectres'][0]:
            mods_summarySpectres = [k for k in content['summarySpectres'][1].keys() if content['summarySpectres'][1][k]]
            self.draw_welch_graphs(db, **kwargs, mods=mods_summarySpectres)
            self.logger.info(f'summarySpectres')

        # pressureTapLocations
        # if db == 'isolated' and content['pressureTapLocations'][0]:
        self.draw_plot_pressure_tap_locations(db, model_size=model_size, alpha=alpha, path_report=path_report)
        self.logger.info(f'pressureTapLocations')
        if db == 'isolated':
            self.draw_plot_model_3d(model_size, path_report)
            self.draw_model_polar(model_size, path_report)

        self.logger.info(f'Отрисовка графиков все')

        # Получение статистики
        self.logger.info(f'Получение статистики')
        # statisticsSensors
        # statistics_sensors = [
        #     угол 0
        #     [
        #         [1 ...],
        #         [2 ...],
        #         [3 ...],
        #         ...
        #     ],
        #     угол 5
        #     [
        #         [1 ...],
        #         [2 ...],
        #         [3 ...],
        #         ...
        #     ],
        # ]
        if content['statisticsSensors'][0]:
            mods_statisticsSensors = [k for k in content['statisticsSensors'][1].keys() if
                                      content['statisticsSensors'][1][k]]
            if mods_statisticsSensors:
                args = [(db, {'mods': mods_statisticsSensors, 'angle': angle}, kwargs) for angle in
                        range(0, angle_border, 5)]
                with ThreadPoolExecutor(max_workers=self._count_threads) as executor:
                    statistics_sensors = list(
                        executor.map(lambda i: self.get_sensor_statistics(i[0], **i[1], **i[2]), args))

        # statisticsSummaryCoefficients
        # statistics_summary_coefficients = [
        #     угол 0
        #     [
        #         [cx ...],
        #         [cy ...],
        #         [cmz ...],
        #     ],
        #     угол 5
        #     [
        #         [cx ...],
        #         [cy ...],
        #         [cmz ...],
        #     ],
        # ]
        if content['summaryCoefficients'][0]:
            # mods_statisticsSummaryCoefficients = [k for k in content['summaryCoefficients'][1].keys()
            #                                       if content['summaryCoefficients'][1][k]]

            mods_statisticsSummaryCoefficients = ['max',
                                                  'mean',
                                                  'min',
                                                  'std',
                                                  'rms',
                                                  'rach',
                                                  'obesP',
                                                  'obesM',
                                                  ]

            args = [(db, {'mods': mods_statisticsSummaryCoefficients, 'angle': angle}, kwargs) for angle in
                    range(0, angle_border, 5)]
            with ThreadPoolExecutor(max_workers=self._count_threads) as executor:
                statistics_summary_coefficients = list(
                    executor.map(lambda i: self.get_summary_coefficients_statistics(i[0], **i[1], **i[2]), args))

        self.logger.info(f'Получение статистики все')

        # Работа с word файлом
        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(14)
        style.font.name = 'Times New Roman'
        section = doc.sections[0]
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        # ширина A4 210 мм высота 297 мм
        fig_width = Mm(165)
        fig_height = Mm(297 / 2 - 55)

        # Шрифт заголовков разного уровня
        head_lvl1 = Pt(20)
        head_lvl2 = Pt(16)
        head_lvl3 = Pt(16)

        counter_plots = 1  # Счетчик графиков для нумерации
        counter_tables = 1  # Счетчик таблиц для нумерации
        counter_head = 1  # Счетчик заголовков
        counter_page_break = 0  # Счетчик для разрыва страницы, каждые 2 рисунка разрыв

        breadth = float(breadth)
        depth = float(depth)
        height = float(height)

        breadth = int(breadth) if breadth.is_integer() else f'{round(breadth, 2):.2f}'
        depth = int(depth) if depth.is_integer() else f'{round(depth, 2):.2f}'
        height = int(height) if height.is_integer() else f'{round(height, 2):.2f}'

        title = doc.add_heading()
        run = title.add_run(f'Отчет по зданию {breadth}x{depth}x{height}')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(24)
        run.bold = True

        for i in ('Параметры ветрового районирования:',
                  f'Ветровой район: {pressure_plot_parameters["wind_region"]}',
                  f'Тип местности: {pressure_plot_parameters["type_area"]}'
                  ):
            doc.add_paragraph().add_run(i)
        if db == 'isolated':
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(f'{path_report}\\Модель\\Модель трехмерная.png', width=fig_width / 2)
            run.add_picture(f'{path_report}\\Модель\\Модель в полярной системе.png', width=fig_width / 2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(f'Рисунок {counter_plots}. '
                                  f'Геометрические размеры и система координат направления ветровых потоков')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            counter_plots += 1

        self.logger.info('1. Геометрические размеры здания')

        head = doc.add_heading()
        run = head.add_run(f'{counter_head}. Геометрические размеры здания')
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = head_lvl1

        counter_head += 1

        header_model = (
            'Геометрический размер',
            'Значение, м'
        )

        table_model = doc.add_table(rows=1, cols=len(header_model))
        table_model.style = 'Table Grid'

        head_cells = table_model.rows[0].cells
        for i, name in enumerate(header_model):
            p = head_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell = p.add_run(name)
            cell.bold = True

        for i, j in zip((breadth, depth, height), ('Ширина:', 'Глубина:', 'Высота:')):
            row_cells = table_model.add_row().cells
            row_cells[0].text = j
            row_cells[1].text = str(i)
        doc.add_page_break()

        # Создание содержания
        p = doc.add_paragraph()
        run = p.add_run('Содержание')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = head_lvl1
        run.bold = True

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)

        doc.add_page_break()

        # if content['pressureTapLocations'][0]:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(f'{path_report}\\Модель\\Развертка модели.png', width=fig_width)
        run.add_text(f'Рисунок {counter_plots}. Система датчиков мониторинга')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        counter_plots += 1

        doc.add_page_break()

        self.logger.info(f'{counter_head}. Статистика по датчикам')
        if content['envelopes'][0]:
            head = doc.add_heading()
            run = head.add_run(f'{counter_head}. Статистика по датчикам')
            run.font.size = head_lvl1
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        counter_head += 1
        for mode in os.listdir(f'{path_report}\\Огибающие'):
            for i2 in os.listdir(f'{path_report}\\Огибающие\\{mode}'):
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(f'{path_report}\\Огибающие\\{mode}\\{i2}', height=fig_height)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                counter_plots += 1

                p = doc.add_paragraph(f'Рисунок {counter_plots}. {i2[:-4]}')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                counter_page_break += 1
                if counter_page_break == 2:
                    doc.add_page_break()
                    counter_page_break = 0

        if counter_page_break != 0:
            doc.add_page_break()

        self.logger.info(f'{counter_head}. Изополя ветровых нагрузок и воздействий')

        if content['isofieldsPressure'][0] or content['isofieldsCoefficients'][0]:
            head = doc.add_heading()
            run = head.add_run(f'{counter_head}. Изополя ветровых нагрузок и воздействий')
            run.font.size = head_lvl1
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.logger.info('3.1 Изополя аэродинамических коэффициентов')
        counter_head_lvl2 = 0
        if content['isofieldsCoefficients'][0]:
            counter_head_lvl2 += 1
            head = doc.add_heading(level=2)
            run = head.add_run(f'{counter_head}.{counter_head_lvl2}. Изополя аэродинамических коэффициентов')
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = head_lvl2

            counter_head_lvl3 = 0

            path_temp = 'Изополя ветровых нагрузок и воздействий'
            for mode in os.listdir(f'{path_report}\\{path_temp}\\Коэффициенты'):
                counter_head_lvl3 += 1
                head = doc.add_heading(level=3)
                run = head.add_run(
                    f'{counter_head}.{counter_head_lvl2}.{counter_head_lvl3}. Изополя {mode_names[mode]} аэродинамических коэффициентов')
                head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = head_lvl3

                for i2 in os.listdir(f'{path_report}\\{path_temp}\\Коэффициенты\\{mode}'):
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(f'{path_report}\\{path_temp}\\Коэффициенты\\{mode}\\{i2}',
                                    height=fig_height - Mm(10))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    counter_plots += 1

                    temp_angle = i2[:i2.rfind(' ')]
                    temp_angle = temp_angle[temp_angle.rfind(' ') + 1:]
                    temp_name = f'Изополя {mode_names[mode]} коэффициентов для угла атаки ветра {temp_angle} º'

                    p = doc.add_paragraph(f'Рисунок {counter_plots}. {temp_name}')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    counter_page_break += 1
                    if counter_page_break == 2:
                        doc.add_page_break()
                        counter_page_break = 0
        if counter_page_break != 0:
            doc.add_page_break()
        self.logger.info('3.2 Изополя давления')
        if content['isofieldsPressure'][0]:
            counter_head_lvl2 += 1

            head = doc.add_heading(level=2)
            run = head.add_run(f'{counter_head}.{counter_head_lvl2}. Изополя давления')
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = head_lvl2

            counter_head_lvl3 = 0

            path_temp = 'Изополя ветровых нагрузок и воздействий'
            for mode in os.listdir(f'{path_report}\\{path_temp}\\Давление'):
                counter_head_lvl3 += 1
                head = doc.add_heading(level=3)
                run = head.add_run(
                    f'{counter_head}.{counter_head_lvl2}.{counter_head_lvl3}. Изополя {mode_names[mode]} давлений')
                head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = head_lvl3

                for i2 in os.listdir(f'{path_report}\\{path_temp}\\Давление\\{mode}'):
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(f'{path_report}\\{path_temp}\\Давление\\{mode}\\{i2}', height=fig_height - Mm(10))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    counter_plots += 1

                    temp_angle = i2[:i2.rfind(' ')]
                    temp_angle = temp_angle[temp_angle.rfind(' ') + 1:]
                    temp_name = f'Изополя {mode_names[mode]} давлений для угла атаки ветра {temp_angle} º'

                    p = doc.add_paragraph(f'Рисунок {counter_plots}. {temp_name}')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    counter_page_break += 1
                    if counter_page_break == 2:
                        doc.add_page_break()
                        counter_page_break = 0

            if counter_page_break != 0:
                doc.add_page_break()

        counter_head += 1

        self.logger.info('4. Статистика по датчикам в табличном виде')

        if content['statisticsSensors'][0]:
            head = doc.add_heading()
            run = head.add_run(f'{counter_head}. Статистика по датчикам в табличном виде ')
            run.font.size = head_lvl1
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER

            counter_head += 1

            header_sensors = ['Датчик']
            header_sensors.extend([id_to_name[mode] for mode in mods_statisticsSensors])

            for angle in range(0, angle_border, 5):
                self.logger.info(f'4. Статистика по датчикам в табличном виде угол {angle}')
                p = doc.add_paragraph()
                p.add_run(
                    f'\nТаблица {counter_tables}. Аэродинамический коэффициент в датчиках для '
                    f'здания {breadth}x{depth}x{height} угол {angle:02}º')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                counter_tables += 1

                table_sensors = doc.add_table(rows=1, cols=len(header_sensors))
                table_sensors.style = 'Table Grid'
                head_cells = table_sensors.rows[0].cells
                for i, name in enumerate(header_sensors):
                    p = head_cells[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = p.add_run(name)
                    cell.bold = True
                    cell.font.size = Pt(8)

                for row in statistics_sensors[angle // 5]:
                    cells = table_sensors.add_row().cells
                    for i, value in enumerate(row):
                        if i == 0:
                            cells[i].text = str(int(value))
                        else:
                            value = int(value) if value.is_integer() else f'{value:.2f}'
                            cells[i].text = str(value)

                        cells[i].paragraphs[0].runs[0].font.size = Pt(12)

            doc.add_page_break()

            del statistics_sensors

        self.logger.info(f'{counter_head}. Суммарные значения аэродинамических коэффициентов')
        if content['summaryCoefficients'][0]:
            head = doc.add_heading()
            run = head.add_run(f'{counter_head}. Графики суммарных аэродинамических коэффициентов')
            run.font.size = head_lvl1
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER

            counter_head += 1
            path_temp = 'Суммарные аэродинамические коэффициенты'

            header_sum = ['Сила']
            header_sum.extend([id_to_name[mode] for mode in mods_statisticsSummaryCoefficients])

            for angle in os.listdir(f'{path_report}\\{path_temp}\\Декартовая система координат'):
                for i2 in os.listdir(f'{path_report}\\{path_temp}\\Декартовая система координат\\{angle}'):
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(f'{path_report}\\{path_temp}\\Декартовая система координат\\{angle}\\{i2}',
                                    height=fig_height)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    counter_plots += 1

                    p = doc.add_paragraph(f'Рисунок {counter_plots}. {i2[:-4]}')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                angle = int(angle)
                table = doc.add_table(rows=1, cols=len(header_sum))
                table.style = 'Table Grid'
                head_cells = table.rows[0].cells

                p = doc.add_paragraph(
                    f'Таблица {counter_tables}. Суммарные аэродинамические коэффициенты '
                    f'для здания {breadth}x{depth}x{height} угол {angle:02}º')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                counter_tables += 1

                for i, name in enumerate(header_sum):
                    p = head_cells[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = p.add_run(name)
                    cell.bold = True
                    cell.font.size = Pt(8)

                for row in statistics_summary_coefficients[angle // 5]:
                    cells = table.add_row().cells
                    for i, value in enumerate(row):
                        cells[i].text = str(value)

                        cells[i].paragraphs[0].runs[0].font.size = Pt(12)

            del statistics_summary_coefficients
            doc.add_page_break()

        # if content['statisticsSummaryCoefficients'][0]:
        #     head = doc.add_heading()
        #     run = head.add_run(f'{counter_head}. Статистика суммарных аэродинамических коэффициентов')
        #     run.font.size = head_lvl1
        #     head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #
        #     counter_head += 1
        #
        #     header_sum = ['Сила']
        #     header_sum.extend([id_to_name[mode] for mode in mods_statisticsSummaryCoefficients])
        #
        #     for angle in range(0, angle_border, 5):
        #         table = doc.add_table(rows=1, cols=len(header_sum))
        #         table.style = 'Table Grid'
        #         head_cells = table.rows[0].cells
        #         for i, name in enumerate(header_sum):
        #             p = head_cells[i].paragraphs[0]
        #             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #
        #             cell = p.add_run(name)
        #             cell.bold = True
        #             cell.font.size = Pt(8)
        #
        #         for row in statistics_summary_coefficients[angle // 5]:
        #             cells = table.add_row().cells
        #             for i, value in enumerate(row):
        #                 cells[i].text = str(value)
        #
        #                 cells[i].paragraphs[0].runs[0].font.size = Pt(12)
        #         doc.add_paragraph().add_run(
        #             f'Таблица {counter_tables}. Суммарные аэродинамические коэффициенты '
        #             f'для здания {breadth}x{depth}x{height} угол {angle:02}º')
        #         counter_tables += 1
        #
        #     del statistics_summary_coefficients
        #     doc.add_page_break()

        if content['polarSummaryCoefficients'][0]:
            head = doc.add_heading()
            run = head.add_run(f'{counter_head}. Cуммарные аэродинамических коэффициенты в полярной системе координат')
            run.font.size = head_lvl1
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER

            counter_head += 1
            path_temp = 'Суммарные аэродинамические коэффициенты'
            for mode in os.listdir(f'{path_report}\\{path_temp}\\Полярная система координат'):
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(f'{path_report}\\{path_temp}\\Полярная система координат\\{mode}')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p = doc.add_paragraph(f'Рисунок {counter_plots}. {mode[:-4]}')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                counter_plots += 1
                doc.add_page_break()

        self.logger.info(f'{counter_head}. Спектры cуммарных значений аэродинамических коэффициентов')

        if content['summarySpectres'][0]:
            head = doc.add_heading()
            run = head.add_run(f'{counter_head}. Спектры cуммарных значений аэродинамических коэффициентов')
            run.font.size = head_lvl1
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER

            counter_head += 1
            path_temp = 'Спектральная плотность мощности'
            for mode in os.listdir(f'{path_report}\\{path_temp}\\Логарифмическая шкала'):
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(f'{path_report}\\{path_temp}\\Логарифмическая шкала\\{mode}')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p = doc.add_paragraph(f'Рисунок {counter_plots}. {mode[:-4]}')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                counter_plots += 1

            doc.add_page_break()

        if db == 'isolated':
            doc.save(f'{path_report}\\Отчет ширина {breadth} глубина {depth} высота {height} альфа {alpha}.docx')
            os.startfile(f'{path_report}\\Отчет ширина {breadth} глубина {depth} высота {height} альфа {alpha}.docx')

        elif db == 'interference':
            doc.save(f'{path_report}\\Отчет ширина {breadth} глубина {depth} высота {height} вариант {case}.docx')
            os.startfile(f'{path_report}\\Отчет ширина {breadth} глубина {depth} высота {height} вариант {case}.docx')

        self.logger.info(f'{time.time() - t1} Затраченное время на создание отчета')

    # Интегрирование по высоте
    def height_integration_cx_cy_cmz_to_csv(self, db, **kwargs):
        COUNT_DOTS = 100
        model_size = kwargs['model_size']
        angle = kwargs['angle']

        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale, scale_factors = get_model_and_scale_factors(*model_size, alpha)

        elif db == 'interference':
            model_scale, scale_factors = get_model_and_scale_factors_interference(*model_size)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients(db, model_name=model_scale, **kwargs)
        coordinates = self.clipboard_obj.get_coordinates(db, model_scale=model_scale, **kwargs)
        model_name = model_scale

        if db == 'isolated':
            breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
            count_sensors_on_middle_row = int(model_name[0]) * 5
            count_sensors_on_side_row = int(model_name[1]) * 5

        elif db == 'interference':
            height = model_scale / 1000
            breadth, depth = 0.07, 0.07
            count_sensors_on_middle = 7
            count_sensors_on_side = 7

        count_sensors_on_model = len(pressure_coefficients[0])

        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle_row + count_sensors_on_side_row))
        count_sensors_on_row = 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)

        angle = int(angle)
        shirina = np.cos(np.deg2rad(angle)) * breadth + np.sin(np.deg2rad(angle)) * depth

        # центры граней
        mid13_x = breadth / 2
        mid24_x = depth / 2

        x1 = coordinates[0]
        x1 = np.reshape(x1, (count_row, -1))
        x1 = np.split(x1, [count_sensors_on_middle_row,
                           count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                           ], axis=1)

        v2 = breadth
        v3 = breadth + depth
        v4 = 2 * breadth + depth
        x1[1] -= v2
        x1[2] -= v3
        x1[3] -= v4

        # mx плечи для каждого сенсора
        mx13 = np.array([
            x1[0] - mid13_x,
            x1[2] - mid13_x,
        ])

        mx24 = np.array([
            x1[1] - mid24_x,
            x1[3] - mid24_x,
        ])

        # Площадь
        s13 = breadth * height
        s24 = depth * height

        x = coordinates[0]
        x = np.reshape(x, (-1, count_sensors_on_row))
        x = np.append(x, [[2 * (breadth + depth)] for _ in range(len(x))], axis=1)
        x = np.insert(x, 0, 0, axis=1)

        y = coordinates[1]
        y = [height for _ in range(count_sensors_on_row)] + y
        y = np.reshape(y, (-1, count_sensors_on_row))
        y = np.append(y, [[0] for _ in range(count_sensors_on_row)])
        y = np.reshape(y, (-1, count_sensors_on_row))

        squares = []
        for y_i in range(count_row):
            for x_i in range(count_sensors_on_row):
                y_t = y[y_i][x_i]
                y_m = y[y_i + 1][x_i]
                y_b = y[y_i + 2][x_i]
                if y_i == 0:
                    dy = y_t - y_m + (y_m - y_b) / 2
                elif y_i == count_row - 1:
                    dy = (y_t - y_m) / 2 + y_m - y_b
                else:
                    dy = (y_t - y_m) / 2 + (y_m - y_b) / 2

                x_l = x[y_i][x_i]
                x_m = x[y_i][x_i + 1]
                x_r = x[y_i][x_i + 2]

                if x_i == 0:
                    dx = x_m - x_l + (x_r - x_m) / 2
                elif x_i == count_sensors_on_row - 1:
                    dx = (x_m - x_l) / 2 + x_r - x_m
                else:
                    dx = (x_m - x_l) / 2 + (x_r - x_m) / 2

                squares.append(dy * dx)
        squares_faces = np.reshape(squares, (count_row, -1))
        squares_faces = np.split(squares_faces, [count_sensors_on_middle_row,
                                                 count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                                                 ], axis=1)

        cx = [
            [] for _ in range(count_row)
        ]
        cy = [
            [] for _ in range(count_row)
        ]
        cmz = [
            [] for _ in range(count_row)
        ]
        for pr in pressure_coefficients:
            pr = np.reshape(pr, (count_row, -1))
            pr = np.split(pr, [count_sensors_on_middle_row,
                               count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                               ], axis=1)

            for row_i in range(count_row):
                faces_x = []
                faces_y = []
                for face in range(4):
                    if face in [0, 2]:
                        faces_x.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s13 / count_row))
                    else:
                        faces_y.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s24 / count_row))

                cx[row_i].append(faces_x[0] - faces_x[1])
                cy[row_i].append(faces_y[0] - faces_y[1])

                # print(mx13[0][row_i])
                # print(pr[0][row_i])
                # print(squares_faces[0][row_i])
                # print(mx13[0][row_i] * pr[0][row_i] * squares_faces[0][row_i])
                t1 = np.sum(mx13[0][row_i] * pr[0][row_i] * squares_faces[0][row_i]) / ((s13 / count_row) * shirina)
                t2 = np.sum(mx24[0][row_i] * pr[1][row_i] * squares_faces[1][row_i]) / ((s24 / count_row) * shirina)
                t3 = np.sum(mx13[1][row_i] * pr[2][row_i] * squares_faces[2][row_i]) / ((s13 / count_row) * shirina)
                t4 = np.sum(mx24[1][row_i] * pr[3][row_i] * squares_faces[3][row_i]) / ((s24 / count_row) * shirina)

                cmz[row_i] = np.append(cmz[row_i], sum([t1, t2, t3, t4]))

        figs = []

        ox = np.linspace(0, 32.768, 32768)
        #
        # cmz1 = np.array(cx)
        #
        # cmz1 = np.sum(cmz1, axis=0)/count_row
        # num_fig = f'Суммарные коэффициенты {123} {123}'
        # fig, ax = plt.subplots(dpi=Plot.dpi, num=num_fig, clear=True)
        #
        # ax.grid()
        # ax.set_ylabel('Суммарные аэродинамические коэффициенты')
        # ax.set_xlabel('Время, с', labelpad=.3)
        #
        # ax.set_xlim(0, 32.768)
        #
        # ax.plot(ox, cmz1, label=f'CMZ')
        #
        # ax.legend(loc='upper right', fontsize=9)
        # ax.set_ylabel('Суммарные аэродинамические коэффициенты')
        # ax.set_xlabel('Время, с', labelpad=.3)
        #
        # figs.append(fig)
        # fig.savefig(f'Интегрирование\\{num_fig}CMZ SUM', bbox_inches='tight')
        #
        # plt.close(fig)

        # from openpyxl import Workbook, load_workbook
        # from openpyxl.styles import Alignment

        # workbook = Workbook()
        # sheet = workbook.active
        # accuracy_values = 2
        #
        # functions = {
        #     'max': lambda d: np.max(d, axis=0).round(accuracy_values),
        #     'mean': lambda d: np.mean(d, axis=0).round(accuracy_values),
        #     'min': lambda d: np.min(d, axis=0).round(accuracy_values),
        #     'std': lambda d: np.std(d, axis=0).round(accuracy_values),
        #     'rms': lambda d: rms(d).round(accuracy_values),
        #     'rach': lambda d: rach(d).round(accuracy_values),
        #     'obesP': lambda d: obes_p(d).round(accuracy_values),
        #     'obesM': lambda d: obes_m(d).round(accuracy_values)
        # }
        #
        # mods = list(functions.keys())
        # head = ['этаж'] + mods
        #

        import pandas as pd

        for i in range(count_row):
            n_arr = np.array([ox.round(5),
                              np.array(cx[i]).round(5),
                              np.array(cy[i]).round(5),
                              np.array(cmz[i]).round(5)
                              ]).T

            df1 = pd.DataFrame(n_arr)

            df1.to_csv(f'Интегрирование\\cx_cy_cmz_115_0_{count_row - i}.csv', index=False, header=False, sep=',')

        # for data, name in zip((cx, cy, cmz), ('CX', 'CY', 'CMZ')):
        #     # sheet.append([name])
        #     # sheet.append(head)
        #     for i, val in enumerate(data):
        # #         data = [count_row - i]
        # #         for mode in mods:
        # #             temp_val = functions[mode](val)
        # #             data.append(int(temp_val) if temp_val.is_integer() else f'{temp_val:.2f}')
        # #         sheet.append(data)
        # #
        # # workbook.save(filename=f'D:\Projects\WindSpectrum\Интегрирование\Статистика {model_name}.xlsx')
        # # workbook.close()
        #
        #
        #
        #
        #         num_fig = f'Суммарные коэффициенты {name} {count_row - i}'
        #         fig, ax = plt.subplots(dpi=Plot.dpi, num=num_fig, clear=True)
        #
        #         ax.grid()
        #         ax.set_ylabel('Суммарные аэродинамические коэффициенты')
        #         ax.set_xlabel('Время, с', labelpad=.3)
        #
        #         ax.set_xlim(0, 32.768)
        #
        #         ax.plot(ox, val, label=f'{name} {count_row - i}')
        #
        #         ax.legend(loc='upper right', fontsize=9)
        #         ax.set_ylabel('Суммарные аэродинамические коэффициенты')
        #         ax.set_xlabel('Время, с', labelpad=.3)
        #
        #         figs.append(fig)
        #         fig.savefig(f'Интегрирование\\{num_fig}', bbox_inches='tight')
        #
        #         plt.close(fig)

        # CX = np.array(cx)
        # CY = np.array(cy)
        #
        # CX = np.mean(cx, axis=0)
        #
        # fig, ax = plt.subplots(dpi=Plot.dpi, num='sfgniksdbgfkjdfs', clear=True)
        #
        # ax.grid()
        # ax.set_ylabel('Суммарные аэродинамические коэффициенты')
        # ax.set_xlabel('Время, с', labelpad=.3)
        #
        # ax.set_xlim(0, 32.768)
        #
        # ax.plot(ox, CX)
        #
        # ax.legend(loc='upper right', fontsize=9)
        # ax.set_ylabel('Суммарные аэродинамические коэффициенты')
        # ax.set_xlabel('Время, с', labelpad=.3)

        return figs

    def height_integration_cx_cy_cmz(self, db, **kwargs):
        COUNT_DOTS = 50
        model_size = kwargs['model_size']
        steps = kwargs['steps']
        plot = kwargs['plot']

        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale, scale_factors = get_model_and_scale_factors(*model_size, alpha)

        elif db == 'interference':
            model_scale, scale_factors = get_model_and_scale_factors_interference(*model_size)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients(db, model_name=model_scale, **kwargs)
        coordinates = self.clipboard_obj.get_coordinates(db, model_scale=model_scale, **kwargs)
        model_name = model_scale

        size_x, size_y, size_z = map(float, model_size)
        x_scale_factor, y_scale_factor, z_scale_factor = scale_factors

        if db == 'isolated':
            breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
            count_sensors_on_middle_row = int(model_name[0]) * 5
            count_sensors_on_side_row = int(model_name[1]) * 5
            if plot == 'spectre':
                angle = int(kwargs['angle'])
                if alpha == '4':
                    speed_sp = speed_sp_b(height)
                    speed_tpu = interp_025_tpu(height)

                elif alpha == '6':
                    speed_sp = speed_sp_a(height)
                    speed_tpu = interp_016_tpu(height)

                l_m = breadth * np.cos(np.deg2rad(angle)) + depth * np.sin(np.deg2rad(angle))
                # print(speed_tpu,l_m)
                sh = lambda f: f * l_m / speed_tpu
        elif db == 'interference':
            height = model_scale / 1000
            breadth, depth = 0.07, 0.07
            count_sensors_on_middle_row = 7
            count_sensors_on_side_row = 7

        count_sensors_on_model = len(pressure_coefficients[0])

        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle_row + count_sensors_on_side_row))
        count_sensors_on_row = 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)

        s13 = breadth * height
        s24 = depth * height

        x, z = np.array(coordinates)

        x = np.reshape(x, (count_row, -1))
        x = np.split(x, [count_sensors_on_middle_row,
                         count_sensors_on_middle_row + count_sensors_on_side_row,
                         2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                         2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                         ], axis=1)

        z = np.reshape(z, (count_row, -1))
        z = np.split(z, [count_sensors_on_middle_row,
                         count_sensors_on_middle_row + count_sensors_on_side_row,
                         2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                         2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                         ], axis=1)

        x_old = [x[i].reshape(1, -1)[0] for i in range(4)]
        z_old = [z[i].reshape(1, -1)[0] * z_scale_factor for i in range(4)]
        for face in range(4):
            # Вычитаем чтобы все координаты по x находились в интервале [0, 1]
            if face == 1:
                x_old[face] -= breadth
                x_old[face] = x_old[face] * y_scale_factor

            elif face == 2:
                x_old[face] -= (breadth + depth)
                x_old[face] = x_old[face] * x_scale_factor

            elif face == 3:
                x_old[face] -= (2 * breadth + depth)
                x_old[face] = x_old[face] * y_scale_factor

            else:
                x_old[face] = x_old[face] * x_scale_factor

        x_marks_border = breadth * x_scale_factor
        y_marks_border = depth * y_scale_factor

        x_dots = [random.uniform(0, x_marks_border) for _ in range(COUNT_DOTS)]
        y_dots = [random.uniform(0, y_marks_border) for _ in range(COUNT_DOTS)]
        z_dots = [[random.uniform(step[0], step[1]) for _ in range(COUNT_DOTS)] for step in steps]
        count_zones = len(steps)

        cx = [[] for _ in range(count_zones)]
        cy = [[] for _ in range(count_zones)]

        squares_x = []
        squares_y = []

        for ind in range(count_zones):
            dz = steps[ind][1] - steps[ind][0]
            squares_x.append(dz * breadth)
            squares_y.append(dz * depth)

        for t_ind, coeff in enumerate(pressure_coefficients):

            # if t_ind == 10000:
            #     break
            print(t_ind)
            coeff = np.reshape(coeff, (count_row, -1))
            coeff = np.split(coeff, [count_sensors_on_middle_row,
                                     count_sensors_on_middle_row + count_sensors_on_side_row,
                                     2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                                     2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                                     ], axis=1)

            faces_x = [[] for _ in range(count_zones)]
            faces_y = [[] for _ in range(count_zones)]

            for face in range(4):
                data_old = coeff[face].reshape(1, -1)[0]

                coords = [[i1, j1] for i1, j1 in zip(x_old[face], z_old[face])]  # Старые координаты
                # Интерполятор полученный на основе имеющихся данных
                interpolator = intp(coords, data_old)

                if face in [0, 2]:
                    for ind, dots in enumerate(z_dots):
                        temp_val = np.mean([float(interpolator([[X, Y]])) for X, Y in zip(x_dots, dots)])
                        faces_x[ind].append(temp_val)
                else:
                    for ind, dots in enumerate(z_dots):
                        temp_val = np.mean([float(interpolator([[X, Y]])) for X, Y in zip(y_dots, dots)])
                        faces_y[ind].append(temp_val)

            for ind in range(count_zones):
                cx[ind].append((faces_x[ind][0] - faces_x[ind][1]))  # * squares_x[ind])
                cy[ind].append((faces_y[ind][0] - faces_y[ind][1]))  # * squares_y[ind])

        figs = []
        labels = []

        # ox = np.linspace(0, 32.768, 10000)
        ox = np.linspace(0, 32.768, 32768)

        for ind in range(count_zones):
            for data, name in zip((cx[ind], cy[ind]), ('CX', 'CY')):

                num_fig = f'Суммарные коэффициенты {"Sh" if plot == "spectre" else ""} {name} {steps[ind][0]} {steps[ind][1]}'
                fig, ax = plt.subplots(dpi=Plot.dpi, num=num_fig, clear=True)

                ax.grid()
                if plot == 'summary':
                    ax.set_ylabel('Суммарные аэродинамические коэффициенты')
                    ax.set_xlabel('Время, с', labelpad=.3)
                    ax.set_xlim(0, 32.768)
                    ax.set_ylabel('Суммарные аэродинамические коэффициенты')
                    ax.plot(ox, data, label=f'{name} {steps[ind][0]} {steps[ind][1]}')

                else:
                    fs = 1000
                    counts = 32768
                    ax.set_xlabel('Sh')
                    ax.set_ylabel('PSD, V**2/Hz')
                    ax.set_xlim([10 ** -1, 2])
                    freq, psd = welch(data, fs=fs, nperseg=int(counts / 5))
                    ax.plot([sh(f) for f in freq], psd, label=f'{name} {steps[ind][0]} {steps[ind][1]}')

                ax.legend(loc='upper right', fontsize=9)

                figs.append(fig)
                labels.append(num_fig)

        return figs, labels

    def height_integration_cx_cy_cmz_floors(self, db, **kwargs):
        model_size = kwargs['model_size']
        angle = kwargs['angle']
        plot = kwargs['plot']

        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale, scale_factors = get_model_and_scale_factors(*model_size, alpha)


        elif db == 'interference':
            model_scale, scale_factors = get_model_and_scale_factors_interference(*model_size)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients(db, model_name=model_scale, **kwargs)
        coordinates = self.clipboard_obj.get_coordinates(db, model_scale=model_scale, **kwargs)
        model_name = model_scale

        if db == 'isolated':
            breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
            count_sensors_on_middle_row = int(model_name[0]) * 5
            count_sensors_on_side_row = int(model_name[1]) * 5

        elif db == 'interference':
            height = model_scale / 1000
            breadth, depth = 0.07, 0.07
            count_sensors_on_middle = 7
            count_sensors_on_side = 7

        count_sensors_on_model = len(pressure_coefficients[0])

        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle_row + count_sensors_on_side_row))
        count_sensors_on_row = 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)

        angle = int(angle)
        shirina = np.cos(np.deg2rad(angle)) * breadth + np.sin(np.deg2rad(angle)) * depth

        # центры граней
        mid13_x = breadth / 2
        mid24_x = depth / 2

        x1 = coordinates[0]
        x1 = np.reshape(x1, (count_row, -1))
        x1 = np.split(x1, [count_sensors_on_middle_row,
                           count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                           ], axis=1)

        v2 = breadth
        v3 = breadth + depth
        v4 = 2 * breadth + depth
        x1[1] -= v2
        x1[2] -= v3
        x1[3] -= v4

        # mx плечи для каждого сенсора
        mx13 = np.array([
            x1[0] - mid13_x,
            x1[2] - mid13_x,
        ])

        mx24 = np.array([
            x1[1] - mid24_x,
            x1[3] - mid24_x,
        ])

        # Площадь
        s13 = breadth * height
        s24 = depth * height

        x = coordinates[0]
        x = np.reshape(x, (-1, count_sensors_on_row))
        x = np.append(x, [[2 * (breadth + depth)] for _ in range(len(x))], axis=1)
        x = np.insert(x, 0, 0, axis=1)

        y = coordinates[1]
        y = [height for _ in range(count_sensors_on_row)] + y
        y = np.reshape(y, (-1, count_sensors_on_row))
        y = np.append(y, [[0] for _ in range(count_sensors_on_row)])
        y = np.reshape(y, (-1, count_sensors_on_row))

        squares = []
        for y_i in range(count_row):
            for x_i in range(count_sensors_on_row):
                y_t = y[y_i][x_i]
                y_m = y[y_i + 1][x_i]
                y_b = y[y_i + 2][x_i]
                if y_i == 0:
                    dy = y_t - y_m + (y_m - y_b) / 2
                elif y_i == count_row - 1:
                    dy = (y_t - y_m) / 2 + y_m - y_b
                else:
                    dy = (y_t - y_m) / 2 + (y_m - y_b) / 2

                x_l = x[y_i][x_i]
                x_m = x[y_i][x_i + 1]
                x_r = x[y_i][x_i + 2]

                if x_i == 0:
                    dx = x_m - x_l + (x_r - x_m) / 2
                elif x_i == count_sensors_on_row - 1:
                    dx = (x_m - x_l) / 2 + x_r - x_m
                else:
                    dx = (x_m - x_l) / 2 + (x_r - x_m) / 2

                squares.append(dy * dx)
        squares_faces = np.reshape(squares, (count_row, -1))
        squares_faces = np.split(squares_faces, [count_sensors_on_middle_row,
                                                 count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                                                 ], axis=1)

        cx = [
            [] for _ in range(count_row)
        ]
        cy = [
            [] for _ in range(count_row)
        ]
        cmz = [
            [] for _ in range(count_row)
        ]
        for pr in pressure_coefficients:
            pr = np.reshape(pr, (count_row, -1))
            pr = np.split(pr, [count_sensors_on_middle_row,
                               count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                               ], axis=1)

            for row_i in range(count_row):
                faces_x = []
                faces_y = []

                for face in range(4):
                    if face in [0, 2]:
                        faces_x.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s13 / count_row))
                    else:
                        faces_y.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s24 / count_row))

                cx[row_i].append(faces_x[0] - faces_x[1])
                cy[row_i].append(faces_y[0] - faces_y[1])

                t1 = np.sum(mx13[0][row_i] * pr[0][row_i] * squares_faces[0][row_i]) / ((s13 / count_row) * shirina)
                t3 = np.sum(mx13[1][row_i] * pr[2][row_i] * squares_faces[2][row_i]) / ((s13 / count_row) * shirina)

                t2 = np.sum(mx24[0][row_i] * pr[1][row_i] * squares_faces[1][row_i]) / ((s24 / count_row) * shirina)
                t4 = np.sum(mx24[1][row_i] * pr[3][row_i] * squares_faces[3][row_i]) / ((s24 / count_row) * shirina)

                cmz[row_i] = np.append(cmz[row_i], sum([t1, t2, t3, t4]))

        if plot == 'spectre':
            angle = int(kwargs['angle'])
            if alpha == '4':
                speed_sp = speed_sp_b(height)
                speed_tpu = interp_025_tpu(height)

            elif alpha == '6':
                speed_sp = speed_sp_a(height)
                speed_tpu = interp_016_tpu(height)

            l_m = breadth * np.cos(np.deg2rad(angle)) + depth * np.sin(np.deg2rad(angle))
            # print(speed_tpu,l_m)
            sh = lambda f: f * l_m / speed_tpu

            return self.plot_integrated_summary_sh(sh, cx, cy, cmz)
        else:
            return self.plot_integrated_summary(cx, cy, cmz)

    def get_height_integration_cx_cy_cmz_floors(self, db, **kwargs):
        model_size = kwargs['model_size']
        angle = kwargs['angle']
        plot = kwargs['plot']

        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale, scale_factors = get_model_and_scale_factors(*model_size, alpha)


        elif db == 'interference':
            model_scale, scale_factors = get_model_and_scale_factors_interference(*model_size)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients(db, model_name=model_scale, **kwargs)
        coordinates = self.clipboard_obj.get_coordinates(db, model_scale=model_scale, **kwargs)
        model_name = model_scale

        if db == 'isolated':
            breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
            count_sensors_on_middle_row = int(model_name[0]) * 5
            count_sensors_on_side_row = int(model_name[1]) * 5

        elif db == 'interference':
            height = model_scale / 1000
            breadth, depth = 0.07, 0.07
            count_sensors_on_middle = 7
            count_sensors_on_side = 7

        count_sensors_on_model = len(pressure_coefficients[0])

        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle_row + count_sensors_on_side_row))
        count_sensors_on_row = 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)

        angle = int(angle)
        shirina = np.cos(np.deg2rad(angle)) * breadth + np.sin(np.deg2rad(angle)) * depth

        # центры граней
        mid13_x = breadth / 2
        mid24_x = depth / 2

        x1 = coordinates[0]
        x1 = np.reshape(x1, (count_row, -1))
        x1 = np.split(x1, [count_sensors_on_middle_row,
                           count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                           ], axis=1)

        v2 = breadth
        v3 = breadth + depth
        v4 = 2 * breadth + depth
        x1[1] -= v2
        x1[2] -= v3
        x1[3] -= v4

        # mx плечи для каждого сенсора
        mx13 = np.array([
            x1[0] - mid13_x,
            x1[2] - mid13_x,
        ])

        mx24 = np.array([
            x1[1] - mid24_x,
            x1[3] - mid24_x,
        ])

        # Площадь
        s13 = breadth * height
        s24 = depth * height

        x = coordinates[0]
        x = np.reshape(x, (-1, count_sensors_on_row))
        x = np.append(x, [[2 * (breadth + depth)] for _ in range(len(x))], axis=1)
        x = np.insert(x, 0, 0, axis=1)

        y = coordinates[1]
        y = [height for _ in range(count_sensors_on_row)] + y
        y = np.reshape(y, (-1, count_sensors_on_row))
        y = np.append(y, [[0] for _ in range(count_sensors_on_row)])
        y = np.reshape(y, (-1, count_sensors_on_row))

        squares = []
        for y_i in range(count_row):
            for x_i in range(count_sensors_on_row):
                y_t = y[y_i][x_i]
                y_m = y[y_i + 1][x_i]
                y_b = y[y_i + 2][x_i]
                if y_i == 0:
                    dy = y_t - y_m + (y_m - y_b) / 2
                elif y_i == count_row - 1:
                    dy = (y_t - y_m) / 2 + y_m - y_b
                else:
                    dy = (y_t - y_m) / 2 + (y_m - y_b) / 2

                x_l = x[y_i][x_i]
                x_m = x[y_i][x_i + 1]
                x_r = x[y_i][x_i + 2]

                if x_i == 0:
                    dx = x_m - x_l + (x_r - x_m) / 2
                elif x_i == count_sensors_on_row - 1:
                    dx = (x_m - x_l) / 2 + x_r - x_m
                else:
                    dx = (x_m - x_l) / 2 + (x_r - x_m) / 2

                squares.append(dy * dx)
        squares_faces = np.reshape(squares, (count_row, -1))
        squares_faces = np.split(squares_faces, [count_sensors_on_middle_row,
                                                 count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                                                 ], axis=1)

        cx = [
            [] for _ in range(count_row)
        ]
        cy = [
            [] for _ in range(count_row)
        ]
        cmz = [
            [] for _ in range(count_row)
        ]
        for pr in pressure_coefficients:
            pr = np.reshape(pr, (count_row, -1))
            pr = np.split(pr, [count_sensors_on_middle_row,
                               count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                               ], axis=1)

            for row_i in range(count_row):
                faces_x = []
                faces_y = []

                for face in range(4):
                    if face in [0, 2]:
                        faces_x.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s13 / count_row))
                    else:
                        faces_y.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s24 / count_row))

                cx[row_i].append(faces_x[0] - faces_x[1])
                cy[row_i].append(faces_y[0] - faces_y[1])

                t1 = np.sum(mx13[0][row_i] * pr[0][row_i] * squares_faces[0][row_i]) / ((s13 / count_row) * shirina)
                t3 = np.sum(mx13[1][row_i] * pr[2][row_i] * squares_faces[2][row_i]) / ((s13 / count_row) * shirina)

                t2 = np.sum(mx24[0][row_i] * pr[1][row_i] * squares_faces[1][row_i]) / ((s24 / count_row) * shirina)
                t4 = np.sum(mx24[1][row_i] * pr[3][row_i] * squares_faces[3][row_i]) / ((s24 / count_row) * shirina)

                cmz[row_i] = np.append(cmz[row_i], sum([t1, t2, t3, t4]))

        return cx, cy, cmz

    def height_integration_cx_cy_cmz_floors_to_txt(self, db, **kwargs):
        # добавить кол во этажей после 6
        # в высоту добавить -1 в начало, для времени

        model_size = kwargs['model_size']
        angle = kwargs['angle']

        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale, scale_factors = get_model_and_scale_factors(*model_size, alpha)


        elif db == 'interference':
            model_scale, scale_factors = get_model_and_scale_factors_interference(*model_size)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients(db, model_name=model_scale, **kwargs)
        coordinates = self.clipboard_obj.get_coordinates(db, model_scale=model_scale, **kwargs)
        uh_speed = np.round(float(self.clipboard_obj.get_uh_average_wind_speed(db, alpha, model_scale)), 3)

        model_name = model_scale

        if db == 'isolated':
            breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
            count_sensors_on_middle_row = int(model_name[0]) * 5
            count_sensors_on_side_row = int(model_name[1]) * 5

        elif db == 'interference':
            height = model_scale / 1000
            breadth, depth = 0.07, 0.07
            count_sensors_on_middle = 7
            count_sensors_on_side = 7

        count_sensors_on_model = len(pressure_coefficients[0])

        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle_row + count_sensors_on_side_row))
        count_sensors_on_row = 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)

        angle = int(angle)
        shirina = np.cos(np.deg2rad(angle)) * breadth + np.sin(np.deg2rad(angle)) * depth

        # центры граней
        mid13_x = breadth / 2
        mid24_x = depth / 2

        x1 = coordinates[0]
        x1 = np.reshape(x1, (count_row, -1))
        x1 = np.split(x1, [count_sensors_on_middle_row,
                           count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                           ], axis=1)

        v2 = breadth
        v3 = breadth + depth
        v4 = 2 * breadth + depth
        x1[1] -= v2
        x1[2] -= v3
        x1[3] -= v4

        # mx плечи для каждого сенсора
        mx13 = np.array([
            x1[0] - mid13_x,
            x1[2] - mid13_x,
        ])

        mx24 = np.array([
            x1[1] - mid24_x,
            x1[3] - mid24_x,
        ])

        # Площадь
        s13 = breadth * height
        s24 = depth * height

        x = coordinates[0]
        x = np.reshape(x, (-1, count_sensors_on_row))
        x = np.append(x, [[2 * (breadth + depth)] for _ in range(len(x))], axis=1)
        x = np.insert(x, 0, 0, axis=1)

        y = coordinates[1]
        z_levels = sorted(set(y), reverse=True)
        y = [height for _ in range(count_sensors_on_row)] + y
        y = np.reshape(y, (-1, count_sensors_on_row))
        y = np.append(y, [[0] for _ in range(count_sensors_on_row)])
        y = np.reshape(y, (-1, count_sensors_on_row))

        squares = []
        for y_i in range(count_row):
            for x_i in range(count_sensors_on_row):
                y_t = y[y_i][x_i]
                y_m = y[y_i + 1][x_i]
                y_b = y[y_i + 2][x_i]
                if y_i == 0:
                    dy = y_t - y_m + (y_m - y_b) / 2
                elif y_i == count_row - 1:
                    dy = (y_t - y_m) / 2 + y_m - y_b
                else:
                    dy = (y_t - y_m) / 2 + (y_m - y_b) / 2

                x_l = x[y_i][x_i]
                x_m = x[y_i][x_i + 1]
                x_r = x[y_i][x_i + 2]

                if x_i == 0:
                    dx = x_m - x_l + (x_r - x_m) / 2
                elif x_i == count_sensors_on_row - 1:
                    dx = (x_m - x_l) / 2 + x_r - x_m
                else:
                    dx = (x_m - x_l) / 2 + (x_r - x_m) / 2

                squares.append(dy * dx)
        squares_faces = np.reshape(squares, (count_row, -1))
        squares_faces = np.split(squares_faces, [count_sensors_on_middle_row,
                                                 count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                                                 ], axis=1)

        cx = [
            [] for _ in range(count_row)
        ]
        cy = [
            [] for _ in range(count_row)
        ]
        cmz = [
            [] for _ in range(count_row)
        ]
        for pr in pressure_coefficients:
            pr = np.reshape(pr, (count_row, -1))
            pr = np.split(pr, [count_sensors_on_middle_row,
                               count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                               ], axis=1)

            for row_i in range(count_row):
                faces_x = []
                faces_y = []

                for face in range(4):
                    if face in [0, 2]:
                        faces_x.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s13 / count_row))
                    else:
                        faces_y.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s24 / count_row))

                cx[row_i].append(faces_x[0] - faces_x[1])
                cy[row_i].append(faces_y[0] - faces_y[1])

                t1 = np.sum(mx13[0][row_i] * pr[0][row_i] * squares_faces[0][row_i]) / ((s13 / count_row) * shirina)
                t3 = np.sum(mx13[1][row_i] * pr[2][row_i] * squares_faces[2][row_i]) / ((s13 / count_row) * shirina)

                t2 = np.sum(mx24[0][row_i] * pr[1][row_i] * squares_faces[1][row_i]) / ((s24 / count_row) * shirina)
                t4 = np.sum(mx24[1][row_i] * pr[3][row_i] * squares_faces[3][row_i]) / ((s24 / count_row) * shirina)

                cmz[row_i] = np.append(cmz[row_i], sum([t1, t2, t3, t4]))

        cx = (np.array(list(reversed(cx))) / count_row).round(5)
        cy = (np.array(list(reversed(cy))) / count_row).round(5)
        cmz = (np.array(list(reversed(cmz))) / count_row).round(5)

        # fig, ax = plt.subplots(dpi=Plot.dpi, num='snfgkjdsnfkjsdnf', clear=True)
        # ax.plot(list(range(32768)), np.sum(cx, axis=0))

        angle = int(kwargs['angle'])

        time = np.linspace(0, 32.768, 32768).round(5)

        path = r'D:\Projects\WindSpectrum\Интегрирование'
        if not os.path.exists(f'{path}\\{model_scale}_{alpha}'):
            os.mkdir(f'{path}\\{model_scale}_{alpha}')
        path = f'{path}\\{model_scale}_{alpha}'

        if alpha == '4':
            speed_2_3 = np.round(interp_025_real_tpu([height])[0], 3)
        elif alpha == '6':
            speed_2_3 = np.round(interp_016_real_tpu([height])[0], 3)

        file_name = f"{path}\\{model_scale}_{alpha}_{angle}.txt"
        f = open(file_name, 'w')
        f.close()

        f = open(file_name, 'ab')
        alpha_temp = '0.25' if alpha == '4' else '0.16'
        f.write(
            f'{model_scale} Вариант модели\n{breadth}, {depth}, {height} м размеры модели b d h\n'
            f'{alpha_temp} альфа\n{angle} угол\n{uh_speed} Uh скорость на высоте H\n'
            f'{speed_2_3} скорость на высоте 2/3 H\n'
            f'{count_row} количество этажей\n'.encode())

        temp_name_str = 'time, '

        # for ind in range(1, count_row + 1):
        #     temp_name_str += f'cx{count_row - ind}, cy{count_row - ind}, cmz{count_row - ind}, '

        for ind in range(1, count_row + 1):
            temp_name_str += f'cx{ind}, cy{ind}, cmz{ind}, '

        temp_name_str = temp_name_str + 'cxsum, cysum, cmzsum\n'
        f.write(temp_name_str.encode())

        # enumerate_str = ', '.join(map(str, reversed(range(count_row * 3 + 1 + 3)))) + '\n'
        enumerate_str = ', '.join(map(str, range(count_row * 3 + 1 + 3))) + '\n'
        f.write(enumerate_str.encode())

        temp_lvl_str = '-1, '
        for z in reversed(z_levels):
            z /= height
            temp_lvl_str += f'{z}, {z}, {z}, '
        temp_lvl_str += '1, 1, 1\n'

        # h_str = ', '.join(map(str, map(lambda x: x / height, z_levels))) + '\n'
        f.write(temp_lvl_str.encode())

        data_to_txt = np.array([time])

        for ind in range(count_row):
            data_to_txt = np.append(data_to_txt, [cx[ind]], axis=0)
            data_to_txt = np.append(data_to_txt, [cy[ind]], axis=0)
            data_to_txt = np.append(data_to_txt, [cmz[ind]], axis=0)

        data_to_txt = np.append(data_to_txt, [np.sum(cx, axis=0).round(5)], axis=0)
        data_to_txt = np.append(data_to_txt, [np.sum(cy, axis=0).round(5)], axis=0)
        data_to_txt = np.append(data_to_txt, [np.sum(cmz, axis=0).round(5)], axis=0)

        np.savetxt(f, data_to_txt.T, newline="\n", delimiter=',', fmt='%.5f')

        f.close()

    def height_integration_cx_cy_cmz_floors_to_txt_inr(self, db, **kwargs):
        model_size = kwargs['model_size']
        angle = kwargs['angle']
        case = kwargs['case']

        model_scale, scale_factors = get_model_and_scale_factors_interference(*model_size)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients(db, model_name=model_scale, **kwargs)
        coordinates = self.clipboard_obj.get_coordinates(db, model_scale=model_scale, **kwargs)

        model_name = model_scale

        height = model_scale / 1000
        breadth, depth = 0.07, 0.07
        count_sensors_on_middle_row = 7
        count_sensors_on_side_row = 7

        count_sensors_on_model = len(pressure_coefficients[0])

        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle_row + count_sensors_on_side_row))
        count_sensors_on_row = 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)

        angle = int(angle)
        shirina = np.cos(np.deg2rad(angle)) * breadth + np.sin(np.deg2rad(angle)) * depth

        # центры граней
        mid13_x = breadth / 2
        mid24_x = depth / 2

        x1 = coordinates[0]
        x1 = np.reshape(x1, (count_row, -1))
        x1 = np.split(x1, [count_sensors_on_middle_row,
                           count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                           ], axis=1)

        v2 = breadth
        v3 = breadth + depth
        v4 = 2 * breadth + depth
        x1[1] -= v2
        x1[2] -= v3
        x1[3] -= v4

        # mx плечи для каждого сенсора
        mx13 = np.array([
            x1[0] - mid13_x,
            x1[2] - mid13_x,
        ])

        mx24 = np.array([
            x1[1] - mid24_x,
            x1[3] - mid24_x,
        ])

        # Площадь
        s13 = breadth * height
        s24 = depth * height

        x = coordinates[0]
        x = np.reshape(x, (-1, count_sensors_on_row))
        x = np.append(x, [[2 * (breadth + depth)] for _ in range(len(x))], axis=1)
        x = np.insert(x, 0, 0, axis=1)

        y = coordinates[1]
        z_levels = sorted(set(y), reverse=True)
        y = [height for _ in range(count_sensors_on_row)] + y
        y = np.reshape(y, (-1, count_sensors_on_row))
        y = np.append(y, [[0] for _ in range(count_sensors_on_row)])
        y = np.reshape(y, (-1, count_sensors_on_row))

        squares = []
        for y_i in range(count_row):
            for x_i in range(count_sensors_on_row):
                y_t = y[y_i][x_i]
                y_m = y[y_i + 1][x_i]
                y_b = y[y_i + 2][x_i]
                if y_i == 0:
                    dy = y_t - y_m + (y_m - y_b) / 2
                elif y_i == count_row - 1:
                    dy = (y_t - y_m) / 2 + y_m - y_b
                else:
                    dy = (y_t - y_m) / 2 + (y_m - y_b) / 2

                x_l = x[y_i][x_i]
                x_m = x[y_i][x_i + 1]
                x_r = x[y_i][x_i + 2]

                if x_i == 0:
                    dx = x_m - x_l + (x_r - x_m) / 2
                elif x_i == count_sensors_on_row - 1:
                    dx = (x_m - x_l) / 2 + x_r - x_m
                else:
                    dx = (x_m - x_l) / 2 + (x_r - x_m) / 2

                squares.append(dy * dx)
        squares_faces = np.reshape(squares, (count_row, -1))
        squares_faces = np.split(squares_faces, [count_sensors_on_middle_row,
                                                 count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                                                 ], axis=1)

        cx = [
            [] for _ in range(count_row)
        ]
        cy = [
            [] for _ in range(count_row)
        ]
        cmz = [
            [] for _ in range(count_row)
        ]
        for pr in pressure_coefficients:
            pr = np.reshape(pr, (count_row, -1))
            pr = np.split(pr, [count_sensors_on_middle_row,
                               count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                               ], axis=1)

            for row_i in range(count_row):
                faces_x = []
                faces_y = []

                for face in range(4):
                    if face in [0, 2]:
                        faces_x.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s13 / count_row))
                    else:
                        faces_y.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s24 / count_row))

                cx[row_i].append(faces_x[0] - faces_x[1])
                cy[row_i].append(faces_y[0] - faces_y[1])

                t1 = np.sum(mx13[0][row_i] * pr[0][row_i] * squares_faces[0][row_i]) / ((s13 / count_row) * shirina)
                t3 = np.sum(mx13[1][row_i] * pr[2][row_i] * squares_faces[2][row_i]) / ((s13 / count_row) * shirina)

                t2 = np.sum(mx24[0][row_i] * pr[1][row_i] * squares_faces[1][row_i]) / ((s24 / count_row) * shirina)
                t4 = np.sum(mx24[1][row_i] * pr[3][row_i] * squares_faces[3][row_i]) / ((s24 / count_row) * shirina)

                cmz[row_i] = np.append(cmz[row_i], sum([t1, t2, t3, t4]))

        cx = (np.array(list(reversed(cx))) / count_row).round(5)
        cy = (np.array(list(reversed(cy))) / count_row).round(5)
        cmz = (np.array(list(reversed(cmz))) / count_row).round(5)

        # fig, ax = plt.subplots(dpi=Plot.dpi, num='snfgkjdsnfkjsdnf', clear=True)
        # ax.plot(list(range(32768)), np.sum(cx, axis=0))

        angle = int(kwargs['angle'])

        time = np.linspace(0, 7.5, 5858).round(5)

        path = r'D:\Projects\WindSpectrum\Интегрирование'
        if not os.path.exists(path):
            os.mkdir(path)

        path = os.path.join(path, 'Интерференция')
        if not os.path.exists(path):
            os.mkdir(path)

        path = os.path.join(path, f'{model_scale}_{case}')
        if not os.path.exists(path):
            os.mkdir(path)

        file_name = f"{path}\\{model_scale}_{case}_{angle}.txt"
        f = open(file_name, 'w')
        f.close()

        f = open(file_name, 'ab')
        f.write(
            f'{model_scale} Вариант модели\n{breadth}, {depth}, {height} м размеры модели b d h\n'
            f'{angle} угол\n8.2 Uh средняя скорость на высоте H\n'
            f'{count_row} количество этажей\n'.encode())

        temp_name_str = 'time, '

        for ind in range(1, count_row + 1):
            temp_name_str += f'cx{ind}, cy{ind}, cmz{ind}, '

        temp_name_str = temp_name_str + 'cxsum, cysum, cmzsum\n'
        f.write(temp_name_str.encode())

        enumerate_str = ', '.join(map(str, range(count_row * 3 + 1 + 3))) + '\n'
        f.write(enumerate_str.encode())

        temp_lvl_str = '-1, '
        for z in reversed(z_levels):
            z /= height
            z = np.round(z, 3)
            temp_lvl_str += f'{z}, {z}, {z}, '
        temp_lvl_str += '1, 1, 1\n'

        # h_str = ', '.join(map(str, map(lambda x: x / height, z_levels))) + '\n'
        f.write(temp_lvl_str.encode())

        data_to_txt = np.array([time])

        for ind in range(count_row):
            data_to_txt = np.append(data_to_txt, [cx[ind]], axis=0)
            data_to_txt = np.append(data_to_txt, [cy[ind]], axis=0)
            data_to_txt = np.append(data_to_txt, [cmz[ind]], axis=0)

        data_to_txt = np.append(data_to_txt, [np.sum(cx, axis=0).round(5)], axis=0)
        data_to_txt = np.append(data_to_txt, [np.sum(cy, axis=0).round(5)], axis=0)
        data_to_txt = np.append(data_to_txt, [np.sum(cmz, axis=0).round(5)], axis=0)

        np.savetxt(f, data_to_txt.T, newline="\n", delimiter=',', fmt='%.5f')

        f.close()

    def plot_integrated_summary(self, *data_in):
        figs = []
        labels = []
        count_row = len(data_in[0])
        ox = np.linspace(0, 32.768, 32768)

        for data, name in zip(data_in, ('CX', 'CY', 'CMZ')):
            for i, val in enumerate(data):
                num_fig = f'Суммарные коэффициенты {name} {count_row - i}'
                fig, ax = plt.subplots(dpi=Plot.dpi, num=num_fig, clear=True)

                ax.grid()

                ax.set_ylabel('Суммарные аэродинамические коэффициенты')
                ax.set_xlabel('Время, с', labelpad=.3)
                ax.set_xlim(0, 32.768)
                ax.set_ylabel('Суммарные аэродинамические коэффициенты')
                ax.plot(ox, val, label=f'{name} {count_row - i}')

                ax.legend(loc='upper right', fontsize=9)

                figs.append(fig)
                labels.append(num_fig)

        return figs, labels

    def plot_integrated_summary_sh(self, sh, *data_in):
        figs = []
        labels = []
        count_row = len(data_in[0])

        fs = 1000
        counts = 32768

        for data, name in zip(data_in, ('CX', 'CY', 'CMZ')):
            for i, val in enumerate(data):
                num_fig = f'Суммарные коэффициенты Sh {name} {count_row - i}'
                fig, ax = plt.subplots(dpi=Plot.dpi, num=num_fig, clear=True)

                ax.grid()

                ax.set_xlabel('Sh')
                ax.set_ylabel('PSD, V**2/Hz')
                ax.set_xlim([10 ** -1, 2])
                freq, psd = welch(val, fs=fs, nperseg=int(counts / 5))

                ax.plot([sh(f) for f in freq], psd, label=f'{name} {count_row - i}')

                ax.legend(loc='upper right', fontsize=9)

                figs.append(fig)
                labels.append(num_fig)

        return figs, labels

    def sh_floors(self, db='isolated', **kwargs):
        COUNT_DOTS = 100
        model_size = kwargs['model_size']
        angle = int(kwargs['angle'])

        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale, scale_factors = get_model_and_scale_factors(*model_size, alpha)

        elif db == 'interference':
            model_scale, scale_factors = get_model_and_scale_factors_interference(*model_size)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients(db, model_name=model_scale, **kwargs)
        coordinates = self.clipboard_obj.get_coordinates(db, model_scale=model_scale, **kwargs)
        model_name = model_scale

        if db == 'isolated':
            breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
            count_sensors_on_middle_row = int(model_name[0]) * 5
            count_sensors_on_side_row = int(model_name[1]) * 5

            if alpha == '4':
                speed_sp = speed_sp_b(height)
                speed_tpu = interp_025_tpu(height)

            elif alpha == '6':
                speed_sp = speed_sp_a(height)
                speed_tpu = interp_016_tpu(height)

            l_m = breadth * np.cos(np.deg2rad(angle)) + depth * np.sin(np.deg2rad(angle))
            # print(speed_tpu,l_m)
            sh = lambda f: f * l_m / speed_tpu
            # sh = lambda f: f * l_m / speed_sp

        elif db == 'interference':
            height = model_scale / 1000
            breadth, depth = 0.07, 0.07
            count_sensors_on_middle = 7
            count_sensors_on_side = 7

        count_sensors_on_model = len(pressure_coefficients[0])

        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle_row + count_sensors_on_side_row))
        count_sensors_on_row = 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)

        angle = int(angle)
        shirina = np.cos(np.deg2rad(angle)) * breadth + np.sin(np.deg2rad(angle)) * depth

        # центры граней
        mid13_x = breadth / 2
        mid24_x = depth / 2

        x1 = coordinates[0]
        x1 = np.reshape(x1, (count_row, -1))
        x1 = np.split(x1, [count_sensors_on_middle_row,
                           count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                           2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                           ], axis=1)

        v2 = breadth
        v3 = breadth + depth
        v4 = 2 * breadth + depth
        x1[1] -= v2
        x1[2] -= v3
        x1[3] -= v4

        # mx плечи для каждого сенсора
        mx13 = np.array([
            x1[0] - mid13_x,
            x1[2] - mid13_x,
        ])

        mx24 = np.array([
            x1[1] - mid24_x,
            x1[3] - mid24_x,
        ])

        # Площадь
        s13 = breadth * height
        s24 = depth * height

        x = coordinates[0]
        x = np.reshape(x, (-1, count_sensors_on_row))
        x = np.append(x, [[2 * (breadth + depth)] for _ in range(len(x))], axis=1)
        x = np.insert(x, 0, 0, axis=1)

        y = coordinates[1]
        y = [height for _ in range(count_sensors_on_row)] + y
        y = np.reshape(y, (-1, count_sensors_on_row))
        y = np.append(y, [[0] for _ in range(count_sensors_on_row)])
        y = np.reshape(y, (-1, count_sensors_on_row))

        squares = []
        for y_i in range(count_row):
            for x_i in range(count_sensors_on_row):
                y_t = y[y_i][x_i]
                y_m = y[y_i + 1][x_i]
                y_b = y[y_i + 2][x_i]
                if y_i == 0:
                    dy = y_t - y_m + (y_m - y_b) / 2
                elif y_i == count_row - 1:
                    dy = (y_t - y_m) / 2 + y_m - y_b
                else:
                    dy = (y_t - y_m) / 2 + (y_m - y_b) / 2

                x_l = x[y_i][x_i]
                x_m = x[y_i][x_i + 1]
                x_r = x[y_i][x_i + 2]

                if x_i == 0:
                    dx = x_m - x_l + (x_r - x_m) / 2
                elif x_i == count_sensors_on_row - 1:
                    dx = (x_m - x_l) / 2 + x_r - x_m
                else:
                    dx = (x_m - x_l) / 2 + (x_r - x_m) / 2

                squares.append(dy * dx)
        squares_faces = np.reshape(squares, (count_row, -1))
        squares_faces = np.split(squares_faces, [count_sensors_on_middle_row,
                                                 count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                                                 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                                                 ], axis=1)

        cx = [
            [] for _ in range(count_row)
        ]
        cy = [
            [] for _ in range(count_row)
        ]
        cmz = [
            [] for _ in range(count_row)
        ]
        for pr in pressure_coefficients:
            pr = np.reshape(pr, (count_row, -1))
            pr = np.split(pr, [count_sensors_on_middle_row,
                               count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                               2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                               ], axis=1)

            for row_i in range(count_row):
                faces_x = []
                faces_y = []
                for face in range(4):
                    if face in [0, 2]:
                        faces_x.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s13 / count_row))
                    else:
                        faces_y.append(np.sum(pr[face][row_i] * squares_faces[face][row_i]) / (s24 / count_row))

                cx[row_i].append(faces_x[0] - faces_x[1])
                cy[row_i].append(faces_y[0] - faces_y[1])

                t1 = np.sum(mx13[0][row_i] * pr[0][row_i] * squares_faces[0][row_i]) / ((s13 / count_row) * shirina)
                t2 = np.sum(mx24[0][row_i] * pr[1][row_i] * squares_faces[1][row_i]) / ((s24 / count_row) * shirina)
                t3 = np.sum(mx13[1][row_i] * pr[2][row_i] * squares_faces[2][row_i]) / ((s13 / count_row) * shirina)
                t4 = np.sum(mx24[1][row_i] * pr[3][row_i] * squares_faces[3][row_i]) / ((s24 / count_row) * shirina)

                cmz[row_i] = np.append(cmz[row_i], sum([t1, t2, t3, t4]))

        figs = []

        ox = np.linspace(0, 32.768, 32768)
        fs = 1000
        counts = 32768

        for data, name in zip((cx, cy, cmz), ('CX', 'CY', 'CMZ')):
            for i, val in enumerate(data):
                num_fig = f'Суммарные коэффициенты Sh {name} {count_row - i}'
                fig, ax = plt.subplots(dpi=Plot.dpi, num=num_fig, clear=True)

                ax.grid()
                ax.set_xlabel('Sh')
                ax.set_ylabel('PSD, V**2/Hz')

                ax.set_xlim([10 ** -1, 2])

                freq, psd = welch(val, fs=fs, nperseg=int(counts / 5))
                ax.plot([sh(f) for f in freq], psd, label=f'{name} {count_row - i}')

                # ax.plot(ox, val, label=f'{name} {count_row - i}')

                ax.legend(loc='upper right', fontsize=9)

                figs.append(fig)
                fig.savefig(f'Спектры\\{num_fig}', bbox_inches='tight')

                plt.close(fig)

        return figs

    def get_coeff_for_melbourne(self, db, **kwargs):
        model_size = kwargs['model_size']
        mode = kwargs['mode']

        angle = kwargs['angle']
        if db == 'isolated':
            alpha = kwargs['alpha']
            model_scale, scale_factors = get_model_and_scale_factors(*model_size, alpha)

        elif db == 'interference':
            model_scale, scale_factors = get_model_and_scale_factors_interference(*model_size)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients(db, model_name=model_scale, **kwargs)
        coordinates = self.clipboard_obj.get_coordinates(db, model_scale=model_scale, **kwargs)
        model_name = model_scale
        # Виды изополей
        mods = {
            'mean': lambda coefficients: np.mean(coefficients, axis=0),
            'rms': lambda coefficients: np.array([np.sqrt(i.dot(i) / i.size) for i in coefficients.T]),
            'std': lambda coefficients: np.std(coefficients, axis=0),
            'max': lambda coefficients: np.max(coefficients, axis=0),
            'min': lambda coefficients: np.min(coefficients, axis=0),
        }

        size_x, size_y, size_z = map(float, model_size)
        x_scale_factor, y_scale_factor, z_scale_factor = scale_factors

        pressure_coefficients = mods[mode](pressure_coefficients)
        if db == 'isolated':
            alpha = kwargs['alpha']
            breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
            count_sensors_on_middle = int(model_name[0]) * 5
            count_sensors_on_side = int(model_name[1]) * 5

        count_sensors_on_model = len(pressure_coefficients)
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))

        x, z = np.array(coordinates)

        pressure_coefficients = np.reshape(pressure_coefficients, (count_row, -1))
        pressure_coefficients = np.split(pressure_coefficients, [count_sensors_on_middle,
                                                                 count_sensors_on_middle + count_sensors_on_side,
                                                                 2 * count_sensors_on_middle + count_sensors_on_side,
                                                                 2 * (count_sensors_on_middle + count_sensors_on_side)
                                                                 ], axis=1)
        x = np.reshape(x, (count_row, -1))
        x = np.split(x, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        z = np.reshape(z, (count_row, -1))
        z = np.split(z, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        del pressure_coefficients[4]
        del x[4]
        del z[4]

        # x это тензор со всеми координатами граней по ширине,x1...x4 координаты отдельных граней
        x1, x2, x3, x4 = list(x[0]), list(x[1]), list(x[2]), list(x[3])
        # z это тензор со всеми координатами граней по высоте,z1...z4 координаты отдельных граней
        z1, z2, z3, z4 = list(z[0]), list(z[1]), list(z[2]), list(z[3])

        # Расширение матрицы координат по бокам
        for i in range(len(x1)):
            x1[i] = np.append(np.insert(x1[i], 0, 0), breadth)
            x2[i] = np.append(np.insert(x2[i], 0, breadth), breadth + depth)
            x3[i] = np.append(np.insert(x3[i], 0, breadth + depth), 2 * breadth + depth)
            x4[i] = np.append(np.insert(x4[i], 0, 2 * breadth + depth), 2 * (breadth + depth))

        x1.append(x1[0])
        x2.append(x2[0])
        x3.append(x3[0])
        x4.append(x4[0])

        x1.insert(0, x1[0])
        x2.insert(0, x2[0])
        x3.insert(0, x3[0])
        x4.insert(0, x4[0])

        # Расширение матрицы координат по бокам
        for i in range(len(z1)):
            z1[i] = np.append(np.insert(z1[i], 0, z1[i][0]), z1[i][0])
            z2[i] = np.append(np.insert(z2[i], 0, z2[i][0]), z2[i][0])
            z3[i] = np.append(np.insert(z3[i], 0, z3[i][0]), z3[i][0])
            z4[i] = np.append(np.insert(z4[i], 0, z4[i][0]), z4[i][0])

        z1.append(np.array([0 for _ in range(len(z1[0]))]))
        z2.append(np.array([0 for _ in range(len(z2[0]))]))
        z3.append(np.array([0 for _ in range(len(z3[0]))]))
        z4.append(np.array([0 for _ in range(len(z4[0]))]))

        z1.insert(0, np.array([height for _ in range(len(z1[0]))]))
        z2.insert(0, np.array([height for _ in range(len(z2[0]))]))
        z3.insert(0, np.array([height for _ in range(len(z3[0]))]))
        z4.insert(0, np.array([height for _ in range(len(z4[0]))]))

        # Расширенные координаты для изополей
        z_extended = [np.array(z1), np.array(z2), np.array(z3), np.array(z4)]
        x_extended = [np.array(x1), np.array(x2), np.array(x3), np.array(x4)]

        x_new_list = []
        z_new_list = []

        x_old_list = []
        z_old_list = []

        data_new_list = []

        ft = 0.3048

        x_melbourne_meters = np.array([15, 45, 75, 105, 135]) * ft
        y_melbourne_meters = np.array([10, 30, 50, 70, 90]) * ft
        z_2_3 = [size_z * 2 / 3 for _ in range(5)]
        data_new_out = []
        for i in range(4):
            x_new = x_extended[i].reshape(1, -1)[0]
            x_old = x[i].reshape(1, -1)[0]

            z_new = z_extended[i].reshape(1, -1)[0]
            z_old = z[i].reshape(1, -1)[0]
            # Вычитаем чтобы все координаты по x находились в интервале [0, 1]
            if i == 1:
                x_old -= breadth
                x_new -= breadth
            elif i == 2:
                x_old -= (breadth + depth)
                x_new -= (breadth + depth)
            elif i == 3:
                x_old -= (2 * breadth + depth)
                x_new -= (2 * breadth + depth)

            # Масштабирование координат
            z_new = z_new * z_scale_factor
            z_old = z_old * z_scale_factor

            if i in [0, 2]:
                x_new = x_new * x_scale_factor
                x_old = x_old * x_scale_factor

            else:
                x_new = x_new * y_scale_factor
                x_old = x_old * y_scale_factor

            data_old = pressure_coefficients[i].reshape(1, -1)[0]

            # data_old_integer.append(data_old)
            coords = [[i1, j1] for i1, j1 in zip(x_old, z_old)]  # Старые координаты
            # Интерполятор полученный на основе имеющихся данных
            interpolator = intp(coords, data_old)
            if i in (0, 2):
                data_new = [float(interpolator([[X, Y]])) for X, Y in zip(x_melbourne_meters, z_2_3)]
            else:
                data_new = [float(interpolator([[X, Y]])) for X, Y in zip(y_melbourne_meters, z_2_3)]

            data_new_out.append(list(reversed(data_new)))
            # data_new_out += data_new

        return data_new_out

    def FEA(self, **kwargs):
        steps = kwargs['steps']
        model_size = kwargs['model_size']
        angle = kwargs['angle']
        alpha = kwargs['alpha']

        COUNT_DOTS_X = 5
        COUNT_DOTS_Y = 5
        COUNT_DOTS_Z = 25

        model_scale, scale_factors = get_model_and_scale_factors(*model_size, alpha)

        pressure_coefficients = self.clipboard_obj.get_pressure_coefficients('isolated', alpha=alpha,
                                                                             model_name=model_scale,
                                                                             angle=angle)
        coordinates = self.clipboard_obj.get_coordinates('isolated', alpha=alpha,
                                                         model_scale=model_scale,
                                                         angle=angle)
        model_name = model_scale

        size_x, size_y, size_z = map(float, model_size)
        x_scale_factor, y_scale_factor, z_scale_factor = scale_factors

        breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
        count_sensors_on_middle_row = int(model_name[0]) * 5
        count_sensors_on_side_row = int(model_name[1]) * 5

        count_sensors_on_model = len(pressure_coefficients[0])

        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle_row + count_sensors_on_side_row))
        count_sensors_on_row = 2 * (count_sensors_on_middle_row + count_sensors_on_side_row)

        s13 = breadth * height
        s24 = depth * height

        x, z = np.array(coordinates)

        x = np.reshape(x, (count_row, -1))
        x = np.split(x, [count_sensors_on_middle_row,
                         count_sensors_on_middle_row + count_sensors_on_side_row,
                         2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                         2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                         ], axis=1)

        z = np.reshape(z, (count_row, -1))
        z = np.split(z, [count_sensors_on_middle_row,
                         count_sensors_on_middle_row + count_sensors_on_side_row,
                         2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                         2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                         ], axis=1)

        x_old = [x[i].reshape(1, -1)[0] for i in range(4)]
        z_old = [z[i].reshape(1, -1)[0] * z_scale_factor for i in range(4)]
        for face in range(4):
            # Вычитаем чтобы все координаты по x находились в интервале [0, 1]
            if face == 1:
                x_old[face] -= breadth
                x_old[face] = x_old[face] * y_scale_factor

            elif face == 2:
                x_old[face] -= (breadth + depth)
                x_old[face] = x_old[face] * x_scale_factor

            elif face == 3:
                x_old[face] -= (2 * breadth + depth)
                x_old[face] = x_old[face] * y_scale_factor

            else:
                x_old[face] = x_old[face] * x_scale_factor

        temp_l_x = np.linspace(0, size_x, COUNT_DOTS_X)
        temp_l_y = np.linspace(0, size_y, COUNT_DOTS_Y)

        temp_l_z = [np.linspace(step[0], step[1], COUNT_DOTS_Z) for step in steps]

        temp_x_d, _ = np.meshgrid(temp_l_x, temp_l_z[0])
        temp_y_d, _ = np.meshgrid(temp_l_y, temp_l_z[0])

        zx_dots = [np.meshgrid(temp_l_x, temp_l_z[i])[1].ravel() for i in range(len(steps))]
        zy_dots = [np.meshgrid(temp_l_y, temp_l_z[i])[1].ravel() for i in range(len(steps))]

        x_dots = temp_x_d.ravel()
        y_dots = temp_y_d.ravel()

        count_zones = len(steps)

        cx = [[] for _ in range(count_zones)]
        cy = [[] for _ in range(count_zones)]
        cmz = [[] for _ in range(count_zones)]

        # снизу вверх
        list_dz = [step[1] - step[0] for step in steps]

        squares_x_face = [i * size_x for i in list_dz]
        squares_y_face = [i * size_y for i in list_dz]

        count_dots_x = COUNT_DOTS_X * COUNT_DOTS_Z
        count_dots_y = COUNT_DOTS_Y * COUNT_DOTS_Z

        # снизу вверх
        squares_x_dot = [i / count_dots_x for i in squares_x_face]
        squares_y_dot = [i / count_dots_y for i in squares_y_face]


        angle = int(angle)
        shirina = np.cos(np.deg2rad(angle)) * breadth + np.sin(np.deg2rad(angle)) * depth
        # центры граней
        mid13_x = size_x / 2
        mid24_x = size_y / 2
        TEST = 5

        for t_ind, coeff in enumerate(pressure_coefficients):

            # if t_ind == TEST:
            #     break
            print(t_ind)
            coeff = np.reshape(coeff, (count_row, -1))
            coeff = np.split(coeff, [count_sensors_on_middle_row,
                                     count_sensors_on_middle_row + count_sensors_on_side_row,
                                     2 * count_sensors_on_middle_row + count_sensors_on_side_row,
                                     2 * (count_sensors_on_middle_row + count_sensors_on_side_row)
                                     ], axis=1)

            faces_x = [[] for _ in range(count_zones)]
            faces_y = [[] for _ in range(count_zones)]
            faces_cmz = [[] for _ in range(count_zones)]

            for face in range(4):
                data_old = coeff[face].reshape(1, -1)[0]

                coords = [[i1, j1] for i1, j1 in zip(x_old[face], z_old[face])]
                # Интерполятор полученный на основе имеющихся данных
                interpolator = intp(coords, data_old)

                if face in [0, 2]:
                    for ind, dots in enumerate(zx_dots):
                        int_dots = np.array([float(interpolator([[X, Y]])) for X, Y in zip(x_dots, dots)])
                        cmz_dot = np.sum(
                            int_dots * np.array([X - mid13_x for X in x_dots]) * squares_x_dot[- ind - 1]) / (
                                          squares_x_face[-ind - 1] * shirina)
                        temp_val = np.sum(int_dots)
                        faces_x[ind].append(temp_val * squares_x_dot[- ind - 1])
                        faces_cmz[ind].append(cmz_dot)
                else:
                    for ind, dots in enumerate(zy_dots):
                        int_dots = np.array([float(interpolator([[X, Y]])) for X, Y in zip(y_dots, dots)])
                        cmz_dot = np.sum(
                            int_dots * np.array([X - mid24_x for X in y_dots]) * squares_y_dot[- ind - 1]) / (
                                          squares_y_face[-ind - 1] * shirina)
                        temp_val = np.sum(int_dots)
                        faces_y[ind].append(temp_val * squares_y_dot[- ind - 1])
                        faces_cmz[ind].append(cmz_dot)

            for ind in range(count_zones):
                cx[ind].append((faces_x[ind][0] - faces_x[ind][1]))
                cy[ind].append((faces_y[ind][0] - faces_y[ind][1]))
                cmz[ind].append(sum(faces_cmz[ind]))

        # figs = []
        # labels = []
        #
        # ox = np.linspace(0, 32.768, TEST)
        # # ox = np.linspace(0, 32.768, 32768)
        #
        # for ind in range(count_zones):
        #     for data, name in zip((cx[ind], cy[ind], cmz[ind]), ('CX', 'CY', 'CMz')):
        #         num_fig = f'Суммарные коэффициенты {name} {steps[ind][0]} {steps[ind][1]}'
        #         fig, ax = plt.subplots(dpi=Plot.dpi, num=num_fig, clear=True)
        #
        #         ax.grid()
        #
        #         ax.set_ylabel('Суммарные аэродинамические коэффициенты')
        #         ax.set_xlabel('Время, с', labelpad=.3)
        #         ax.set_xlim(0, 32.768)
        #         ax.set_ylabel('Суммарные аэродинамические коэффициенты')
        #         ax.plot(ox, data, label=f'{name} {steps[ind][0]} {steps[ind][1]}')
        #
        #         ax.legend(loc='upper right', fontsize=9)
        #
        #         figs.append(fig)
        #         labels.append(num_fig)
        #
        # return figs, labels

        po = 1.225
        v = 10

        shirina = np.cos(np.deg2rad(angle)) * size_x + np.sin(np.deg2rad(angle)) * size_y
        S = 2 * list_dz[0] * (size_x + size_y)
        # снизу вверх
        cx = np.array([po * v ** 2 * S * np.array(cx[-t_ind - 1]) / 2 for t_ind, sx in enumerate(squares_x_face)])
        cy = np.array([po * v ** 2 * S * np.array(cy[-t_ind - 1]) / 2 for t_ind, sy in enumerate(squares_y_face)])
        cmz = np.array(
            [po * v ** 2 * S * np.array(cmz[-t_ind - 1]) * shirina / 2 for t_ind, sy in enumerate(squares_y_face)])
        #ox = np.linspace(0, 32.768, TEST)
        ox = np.linspace(0, 32.768, 32768)

        e = np.concatenate((cx.T, cy.T, cmz.T), axis=1)
        file_name = f"FEA {model_scale}_{alpha}_{angle}.txt"
        f = open(file_name, 'w')
        f.close()
        f = open(file_name, 'ab')
        #np.savetxt(f, e, newline="\n", delimiter=' ', fmt='%.5f')
        for i in range(len(ox)):
            ltemp = [ox[i]]
            for j in range(count_zones):
                ltemp.append(cx[j][i])
                ltemp.append(cy[j][i])
                ltemp.append(cmz[j][i])
            np.savetxt(f, [ltemp], newline="\n", delimiter=' ', fmt='%.5f')
        f.close()





if __name__ == '__main__':
    c = Core()
    c.report('6', ('3', '1', '5'))
    # run_proc(c.clipboard_obj.clipboard_dict, '4', ('1', '1', '1'))
    # fig = c.get_plot_summary_coefficients('6', ('0.1', '0.1', '0.1'), '0', 'Cx Cy CMz')
    # utils.utils.open_fig(fig)
    # t1 = time.time()
    # c.preparation_for_report('6', ('1', '4', '8'))
    # c.get_plot_isofields('6', ('1', '1', '1'),'0','mean','isofields_coefficients')
    # c.save_clipboard()
    # c.preparation_for_report('6', ('1', '1', '1'))
    # print(time.time() - t1)
    # print('=========================================')
    # c.report('6', ('2', '2', '2')) alpha: str,
    #                            model_size: Tuple[str, str, str],
    #                            angle: str,
    #                            mode: str,
    #                            type_plot: str):
    # print(os.listdir('D:\\'))
